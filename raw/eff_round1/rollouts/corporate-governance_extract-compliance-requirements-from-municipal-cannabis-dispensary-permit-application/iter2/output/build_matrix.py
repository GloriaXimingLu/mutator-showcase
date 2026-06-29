#!/usr/bin/env python3
"""Build the Cannabis Dispensary Permit Compliance Requirements Matrix (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- helpers ----------
def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right),
                     ('left', left), ('right', right)):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), 'true')
    trPr.append(th)

def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

def cell_text(cell, text, bold=False, size=8, color=None, align=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    no_space(p)
    if align is not None:
        p.alignment = align
    # support simple multi-paragraph via \n
    parts = text.split('\n')
    for i, part in enumerate(parts):
        if i > 0:
            p = cell.add_paragraph()
            no_space(p)
            if align is not None:
                p.alignment = align
        r = p.add_run(part)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = 'Calibri'
        if color:
            r.font.color.rgb = color
    return cell

def add_heading(doc, text, level=1, color=RGBColor(0x1F, 0x3A, 0x5F)):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = color
        r.font.name = 'Calibri'
    return h

def body(doc, text, size=10, bold=False, italic=False, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    if color:
        r.font.color.rgb = color
    return p

def bullet(doc, text, size=10, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.name = 'Calibri'
    return p

# ---------- status color map ----------
STATUS_COLORS = {
    'Compliant': 'C6EFCE',
    'Partial': 'FFEB9C',
    'Non-Compliant': 'FFC7CE',
    'Gap': 'FFD9B3',
    'Discrepancy': 'D9E1F2',
    'Critical': 'F4A6A6',
    'Future': 'E2E2E2',
}
STATUS_TEXT = {
    'Compliant': RGBColor(0x00, 0x61, 0x00),
    'Partial': RGBColor(0x7F, 0x60, 0x00),
    'Non-Compliant': RGBColor(0x9C, 0x00, 0x06),
    'Gap': RGBColor(0x8A, 0x4B, 0x00),
    'Discrepancy': RGBColor(0x1F, 0x3A, 0x5F),
    'Critical': RGBColor(0x9C, 0x00, 0x06),
    'Future': RGBColor(0x40, 0x40, 0x40),
}

# ---------- document setup ----------
doc = Document()
# base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)

sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

# ---------- title block ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('COMPLIANCE REQUIREMENTS MATRIX')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
r.font.name = 'Calibri'
t.paragraph_format.space_after = Pt(2)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('City of Evanston Cannabis Dispensary Permit Application')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
st.paragraph_format.space_after = Pt(2)

st2 = doc.add_paragraph()
st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run('Applicant: Greenleaf Therapeutics LLC  |  Proposed Location: 1420 Sherman Ave., Unit B, Evanston, IL 60201')
r.font.size = Pt(10.5)
r.italic = True
st2.paragraph_format.space_after = Pt(2)

st3 = doc.add_paragraph()
st3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st3.add_run('Review against Evanston Municipal Code Ch. 9, Art. 5 and the City Application Checklist & Submission Guide')
r.font.size = Pt(10)
r.italic = True
st3.paragraph_format.space_after = Pt(10)

# meta table
meta = doc.add_table(rows=2, cols=4)
meta.style = 'Light Grid Accent 1'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Prepared for', 'Cannabis Licensing Coordinator, City of Evanston', 'Applicant entity', 'Greenleaf Therapeutics LLC (IL LLC, formed 3/14/2023)'),
    ('State license', 'ADC-285-0147 (conditional; exp. 9/21/2024)', 'Target submission', 'On or before 5/31/2024 (window 2/1–5/31/2024)'),
]
for i, row in enumerate(meta_data):
    for j, val in enumerate(row):
        c = meta.rows[i].cells[j]
        cell_text(c, val, bold=(j % 2 == 0), size=8.5,
                  color=RGBColor(0x1F, 0x3A, 0x5F) if j % 2 == 0 else None)
        set_cell_margins(c)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- 1. executive summary ----------
add_heading(doc, '1.  Executive Summary', level=1)
body(doc,
     'This matrix assesses the Greenleaf Therapeutics LLC ("Greenleaf") cannabis dispensary permit '
     'application package against the controlling authority — Evanston Municipal Code Chapter 9, Article 5 '
     '("the Code") — and the City\'s Application Checklist & Submission Guide ("the Guide") and permit '
     'application form (Form v3.2). Each requirement is evaluated for status and accompanied by remediation '
     'notes. Where the City\'s own documents (Code, Guide, and Form) conflict, the discrepancy is flagged '
     'because it directly affects the applicant\'s ability to prepare a conforming submission; the Code '
     'controls where the Guide or Form imposes a conflicting or additional standard.')
body(doc,
     'The package is substantively strong on narrative content (business plan, lease, community engagement, '
     'social-equity narrative) but is NOT submission-ready. Five issues are critical and would, individually '
     'or collectively, cause the application to be deemed incomplete or to lose material scoring points: '
     '(1) the odor-mitigation Professional Engineer certification is not yet available because the system '
     'design is unfinished; (2) the second community-engagement session was noticed only 12 days in advance, '
     'short of the 14-day minimum; (3) no bonding-capacity commitment letter has been obtained; (4) the State '
     'conditional license expires 9/21/2024, before the anticipated local permit issuance; and (5) the '
     'security plan is not yet finalized. Numerous supporting exhibits referenced in the package are also not '
     'present and must be assembled before submission.')

# status summary table
body(doc, 'Status summary across the 38 reviewed requirements:', bold=True, space_after=4)
sumtab = doc.add_table(rows=1, cols=4)
sumtab.style = 'Light List Accent 1'
hdr = sumtab.rows[0].cells
for j, h in enumerate(['Status', 'Count', 'Indicates', 'Color']):
    cell_text(hdr[j], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade(hdr[j], '1F3A5F')
summary_rows = [
    ('Critical', 6, 'Application deemed incomplete / material scoring loss if not cured', 'Critical'),
    ('Partial', 11, 'Item present but incomplete or with open sub-elements', 'Partial'),
    ('Discrepancy (City docs)', 6, 'Code/Guide/Form conflict affecting compliance', 'Discrepancy'),
    ('Compliant', 9, 'Meets the controlling requirement', 'Compliant'),
    ('Gap (not provided)', 1, 'Required item entirely absent from the package', 'Gap'),
    ('Future / ongoing', 5, 'Post-issuance or operational obligation (noted for awareness)', 'Future'),
]
for s, cnt, ind, key in summary_rows:
    row = sumtab.add_row().cells
    cell_text(row[0], s, bold=True, size=9, color=STATUS_TEXT[key])
    shade(row[0], STATUS_COLORS[key])
    cell_text(row[1], str(cnt), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row[2], ind, size=9)
    cell_text(row[3], '', size=9)
    shade(row[3], STATUS_COLORS[key])
    for c in row:
        set_cell_margins(c)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
body(doc,
     'Note: Several requirements carry more than one status dimension (e.g., "Compliant (amount) — '
     'Discrepancy (method)" or "Partial — Gap"). Each is counted above by its most-serious dimension and '
     'color-coded accordingly in the matrix; the per-requirement rows below state all applicable dimensions. '
     'In particular, three items (A10, A12, A17) are labeled "Non-Compliant — Critical gap" and are counted '
     'under Critical because they are application-incomplete issues; items marked "Partial — Gap" (where a '
     'required sub-component is missing) are counted under Partial. No item is solely non-compliant without a '
     'critical dimension.', italic=True, size=9, space_after=8)

# ---------- 2. scope & methodology ----------
add_heading(doc, '2.  Scope, Methodology & Authority Hierarchy', level=1)
body(doc,
     'Scope. The review covers all submission-time requirements of Code §9-5-5 and the Guide Master Checklist '
     '(Items 1–20), the principal operational/scoring requirements of Code §§9-5-9 through 9-5-22, and the '
     'pre-issuance and ongoing obligations the applicant must plan for. The application form (a blank '
     'template) is assessed both for the fields the applicant must complete and for errors the form itself '
     'contains relative to the Code.')
body(doc,
     'Authority hierarchy. Where the Code, Guide, and Form conflict, the Code is controlling (it is the '
     'enacted ordinance; the Guide and Form are administrative aids). However, the Cannabis Licensing '
     'Coordinator administers the process, so an applicant should seek written confirmation where a conflict '
     'exists, and should satisfy the more restrictive of the Code or the Guide/Form where feasible to avoid '
     'a completeness dispute. Discrepancies are logged in Section 6.')
body(doc,
     'Status legend. Compliant = meets the controlling requirement. Partial = present but incomplete or with '
     'open sub-elements. Gap = required item not in the package. Non-Compliant = documented failure to meet a '
     'stated requirement. Discrepancy = the Code, Guide, and/or Form conflict on the requirement. Critical = '
     'an issue that, if uncured, renders the application incomplete or causes material scoring loss. '
     'Future/ongoing = a post-issuance or operational obligation noted for awareness.')

# ---------- 3. documents reviewed ----------
add_heading(doc, '3.  Documents Reviewed', level=1)
docs = [
    ('Evanston Municipal Code, Title 9, Ch. 9, Art. 5', 'Controlling ordinance (Ord. 78-O-23, eff. 1/1/2024)'),
    ('Application Checklist & Submission Guide', 'City administrative guide (pub. 1/15/2024)'),
    ('Cannabis Dispensary Permit Application Form (v3.2)', 'Blank template (rev. 1/12/2024) — not yet completed'),
    ('Greenleaf Business Plan — Executive Summary & Operations Plan', 'Applicant narrative (April 2024)'),
    ('Community Engagement Session Documentation', 'Applicant (prepared 4/15/2024)'),
    ('Insurance & Bonding Status Memorandum', 'Applicant (dated 4/22/2024)'),
    ('Summary of Key Lease Terms', 'Hargrove & Sinclair LLP (April 2024)'),
    ('State Conditional Adult-Use Dispensing Organization License ADC-285-0147', 'IDFPR (issued 9/22/2023)'),
]
dt = doc.add_table(rows=1, cols=2)
dt.style = 'Light List Accent 1'
dh = dt.rows[0].cells
cell_text(dh[0], 'Document', bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
cell_text(dh[1], 'Description / role', bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
shade(dh[0], '1F3A5F'); shade(dh[1], '1F3A5F')
for name, desc in docs:
    row = dt.add_row().cells
    cell_text(row[0], name, size=9, bold=True)
    cell_text(row[1], desc, size=9)
    set_cell_margins(row[0]); set_cell_margins(row[1])
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ============================================================
# 4. MAIN MATRIX (landscape)
# ============================================================
doc.add_page_break()
land = doc.add_section(WD_SECTION.NEW_PAGE)
land.orientation = WD_ORIENT.LANDSCAPE
land.page_width = Inches(11)
land.page_height = Inches(8.5)
land.top_margin = Inches(0.6)
land.bottom_margin = Inches(0.6)
land.left_margin = Inches(0.6)
land.right_margin = Inches(0.6)

add_heading(doc, '4.  Compliance Requirements Matrix', level=1)
body(doc,
     'Requirements are grouped: A = Application Package Components (submission-time, Code §9-5-5 / Guide '
     'Items 1–20); B = Operational, Scoring & Process Requirements. "Authority" cites the Code section, Guide '
     'item, and Form section. Status color is applied to the Status column.', size=9.5, space_after=6)

# matrix data: (id, requirement, authority, findings, status_label, status_key, remediation)
MATRIX = [
 # ---- Group A ----
 ("A1", "Completed Application Form (signed & notarized)",
  "Code §9-5-5(B)(1); Guide Item 1; Form §§1–14",
  "The Form (v3.2) in the package is a BLANK template — no fields completed, no authorized signature, no notarization. Supporting narratives are present, but the form itself is unexecuted.",
  "Gap", "Gap",
  "Complete all fields (mark 'N/A' where inapplicable); obtain the authorized signatory's signature (Okonkwo, Managing Member) and notarization (Form §14). Deliver 3 bound hard copies + 1 USB PDF."),

 ("A2", "Application Fee ($5,000) & payment method",
  "Code §9-5-6; Guide Item 2; Form §13",
  "Fee amount ($5,000) is consistent. PAYMENT-METHOD CONFLICT: Code authorizes certified check or electronic funds transfer (EFT); Guide authorizes cashier's check or money order; Form authorizes certified check, cashier's check, or wire transfer. Money order (Guide) is not Code-authorized; EFT (Code) is omitted by the Guide. No evidence of payment in package.",
  "Discrepancy + Gap", "Discrepancy",
  "Confirm the acceptable method with the Coordinator. To be Code-safe, pay by certified check or EFT (Code-authoritative). Record the payment reference in Form §13."),

 ("A3", "State Cannabis License",
  "Code §9-5-5(B)(2), §9-5-11(I); Guide Item 3; Form §4 Q9",
  "Conditional license ADC-285-0147 provided (issued 9/22/2023; expires 9/21/2024); licensee, EIN, location, ownership, and social-equity designation all match. CRITICAL TIMELINE RISK: the conditional license expires 9/21/2024, but the business plan anticipates local permit issuance 'Late Q3 or Early Q4 2024' (Oct–Nov). If the local permit issues after 9/21/2024, the State license will have lapsed and §9-5-11(I) (valid State license at issuance) cannot be satisfied. The 180-day conversion clock starts at local issuance, but the conditional license itself expires first.",
  "Compliant (provided) — Critical risk", "Critical",
  "File a State extension petition (License Condition 4: up to 6-month extension, filed ≥60 days before expiration, i.e., by ~7/23/2024) if the local permit will not issue before 9/21/2024. Expedite the local process. Include the conversion timeline/conditions documentation (Guide Item 3)."),

 ("A4", "Entity formation & organizational documents",
  "Code §9-5-5(B)(3); Guide Item 4; Form §1, §13",
  "Entity info provided (IL LLC, formed 3/14/2023, EIN 88-3921047, ILSOS File LLC-07842319). REGISTERED-AGENT NAME CONFLICT: business plan & lease say 'Northwest Statutory Services'; State license says 'Keystone Statutory Services' (same address). Certificate of Good Standing: Code/Guide require within 60 days; Form §13 requires within 30 days (Form more restrictive) — not in package. IRS EIN verification letter, Articles of Organization, and Operating Agreement: referenced but not in package. Ownership chart provided (BP §2.2).",
  "Partial — Gap + Discrepancy", "Partial",
  "Reconcile the registered-agent name (Northwest vs. Keystone). Obtain a Certificate of Good Standing (within 60 days per Code; 30 days to also satisfy the Form). Obtain the IRS EIN verification letter (CP 575/147C). Include Articles of Organization and the Operating Agreement."),

 ("A5", "Principal officer & owner disclosures",
  "Code §9-5-5(B)(4), §9-5-22; Guide Item 5; Form §2 Q11/Q12, §10 Q14A/B",
  "Two principal officers >5% disclosed: Okonkwo (60%, Managing Member/Agent-in-Charge) and Ramachandran (40%, Silent Member). Individual Disclosure Forms (App. A), government photo ID, proof of residency (within 90 days), background-check authorizations, fingerprint receipts, and criminal-history disclosures are NOT in the package. SSN DISCREPANCY: Code §9-5-5(B)(4) requires each officer's full SSN; Form Q14A requests last 4 digits only. Okonkwo's 2016 expunged misdemeanor is disclosed and is not disqualifying; Ramachandran has no record.",
  "Partial — Gap + Discrepancy", "Partial",
  "Complete Individual Disclosure Forms for both officers; provide photo ID, proof of residency (within 90 days), signed background-check authorizations, fingerprint receipts, and criminal-history disclosures. Reconcile the SSN requirement (Code = full SSN; Form = last 4 only — confirm with Coordinator)."),

 ("A6", "Proof of right to occupy (lease/deed)",
  "Code §9-5-5(B)(5); Guide Item 6; Form §3 Q4",
  "Lease summary provided. Executed 1/8/2024; 7-year initial term (to 1/7/2031) + two 3-year renewals (to 2037). Landlord (Kimball Properties Inc.) expressly consents to cannabis use (Exhibit D). MEASUREMENT-BASIS DISCREPANCY: Code measures remaining term from 'anticipated date of permit issuance'; Guide/Form measure from 'date of application submission.' The lease has 7+ years remaining under either measure — COMPLIANT regardless. Landlord Cannabis Use Acknowledgment & Consent (Exhibit D) referenced as executed.",
  "Compliant — Minor discrepancy", "Compliant",
  "Include the executed lease and Exhibit D (Landlord Consent) in the final package. The measurement-basis discrepancy has no practical impact given the 7-year term."),

 ("A7", "Site plan & floor plan",
  "Code §9-5-5(B)(6); Guide Item 7; Form §3 Q5",
  "Floor-plan breakdown provided (Retail 1,600; Secure Storage 800; Office/Employee 600; Vestibule 200; Total 3,200 sq ft). Code requires the plan prepared by a licensed architect or engineer — NOT confirmed in the package (the survey was by a licensed land surveyor, but the floor-plan preparer is unidentified). Form §3 Q5 requires extensive elements (camera locations/fields of view, POS, secure vestibule, ADA, parking, signage) — described conceptually in the business plan, but the scaled plan is not in the package.",
  "Partial — Gap", "Partial",
  "Provide a scaled site plan and floor plan (min. 1/8\" = 1 ft) prepared/certified by a licensed IL architect or engineer, including all Form §3 Q5 elements (camera locations & fields of view, POS, vestibule/ID station, secure storage, ADA features, parking, signage)."),

 ("A8", "Buffer distance certification",
  "Code §9-5-5(B)(7), §9-5-9(B),(C),(F); Guide Item 8; Form §3 Q6/Q7/Q8",
  "Survey by Meridian Land Surveyors LLC (IPLS #035-004728), dated 12/15/2023 — surveyor properly licensed. Buffer distances (straight-line): Chute Middle School 1,180 ft (>1,000 ✓); Tiny Acorns Daycare 890 ft (>500 ✓); Prairie Wellness Dispensary 2,100 ft (>1,500 ✓) — all exceed minimums. METHODOLOGY CONFLICT: Code §9-5-9(C) measures 'nearest property line to nearest property line' (straight-line); Form §3 Q7 requires 'nearest pedestrian route' and states straight-line is 'not acceptable.' Applicant used straight-line (Code-compliant) but the Form rejects it; the lease summary §6 flags this. SURVEY STALENESS: Form requires the survey within 60 days of submission; it is ~5 months old. ONE-DISPENSARY-PER-BLOCK (Code §9-5-9(F)): not addressed.",
  "Partial — Discrepancy + Gap", "Discrepancy",
  "Reconcile the measurement methodology with the Coordinator (Code straight-line vs. Form pedestrian-route); obtain pedestrian-route measurements (likely still compliant given margins) or confirm straight-line is acceptable. Obtain an updated survey within 60 days of submission. Document one-dispensary-per-block compliance."),

 ("A9", "Security plan",
  "Code §9-5-5(B)(9), §9-5-13; Guide Item 9; Form §5 Q15",
  "Business plan §5.3 describes a security plan designed by Pinnacle Compliance Advisors LLC (Marcus Wren): 24/7 surveillance (1080p, 30-day retention), panic alarm (direct to EPD), secure vestibule, access control, cash handling, employee training. CRITICAL GAP: the standalone security plan is NOT in the package; the insurance memo §4 confirms it is 'still being finalized by Pinnacle… and has not yet been delivered' (it is also a dependency for the cannabis crime/theft underwriting). Code §9-5-13(A) requires camera coverage of exterior/parking/loading zones and adjacent alleyways — the lease notes a rear service alley; confirm coverage.",
  "Critical — Gap (not finalized)", "Critical",
  "Finalize and submit the complete security plan (Pinnacle) addressing all Code §9-5-13 elements, including exterior/parking/loading/alleyway camera coverage, the quarterly panic-alarm testing protocol, and the security consultant's qualifications (Form Q15(h))."),

 ("A10", "Odor mitigation plan (with PE certification)",
  "Code §9-5-5(B)(10), §9-5-13(G); Guide Item 10; Form §5 Q16",
  "CRITICAL. Code, Guide, and Form ALL require a certification signed, sealed, and dated by an IL-licensed Professional Engineer that the odor-mitigation system is adequate to prevent odor migration beyond the property line — AT THE TIME OF APPLICATION. Form §5 Q16 note: 'Applications submitted without the required P.E. certification will be deemed incomplete and will not proceed to scoring.' Business plan §5.4 states HVAC/filtration specs 'will be finalized during the build-out engineering phase'; the CE report §2.3 states the odor plan 'will be prepared by a licensed mechanical engineer and included… in the application submission.' This indicates the PE certification is NOT yet available because the design is unfinished.",
  "Non-Compliant — Critical gap", "Critical",
  "Finalize the HVAC/carbon-filtration design and obtain the signed/sealed PE certification BEFORE submission. Without it the application is deemed incomplete and will not proceed to scoring. This is a hard submission-time requirement (not deferrable to pre-issuance)."),

 ("A11", "Business plan",
  "Code §9-5-5(B)(8); Guide Item 11; Form §8 Q25",
  "Comprehensive business plan provided: executive summary, market analysis, operations plan, staffing model (12 FTEs), 3-year (actually 5-year) financial projections, inventory/seed-to-sale (COMPASS), customer flow, training, and an organizational chart (Appendix A). Satisfies the Code elements (market analysis, ≥3-year projections, staffing, management structure, operational plan).",
  "Compliant", "Compliant",
  "None. (Minor: provide an explicit FT/PT breakdown for Form §5 Q13; ensure the org chart is attached as Appendix A.)"),

 ("A12", "Community engagement documentation",
  "Code §9-5-8; Guide Item 12; Form §7 Q22",
  "Two sessions held (3/12 & 4/9/2024) at the Evanston Public Library, both evening (6:30 PM) — satisfies Code §9-5-8(D). Attendance sheets provided (Exhibits A & B) with names/addresses. Published notices provided (Exhibit C). CRITICAL — Session 2 notice timing: Session 2 on 4/9/2024; notice published 3/28/2024 = 12 days; Code §9-5-8(B) and the Guide require ≥14 calendar days. The report itself states 'Days of Advance Notice: 12 days.' NON-COMPLIANT. (Session 1: 2/25→3/12 = 16 days, compliant.) NOTICE CONTENT: Code §9-5-8(C)(5) requires a telephone number AND email; the notices give only email — no telephone. WEBSITE POSTING: Code §9-5-8(B)(2),(E)(2) requires City-website posting ≥14 days prior with evidence; claimed posting on the same dates (so Session 2 website posting is also 12 days — non-compliant); confirmations referenced but not in package; 3-business-day advance submission to the Coordinator not documented. PUBLISHER'S AFFIDAVIT (Code §9-5-8(E)(1)) and written public comments (§9-5-8(E)(5)) not included.",
  "Non-Compliant — Critical + gaps", "Critical",
  "CRITICAL — Per Code §9-5-8(F), the deficient Session 2 notice renders the application incomplete; hold an ADDITIONAL community-engagement session with full ≥14-day newspaper + website notice (no waiver of the 2-session minimum). Add a telephone number to all notices; state adult-use AND medical. Obtain publisher's affidavits; document City-website posting confirmations and the 3-business-day advance notice submission to the Coordinator; include any written public comments."),

 ("A13", "Community benefits plan",
  "Code §9-5-5(B)(12), §9-5-15; Guide Item 13; Form §8 Q28",
  "Business plan §7.4 commits to the greater of $25,000 or 1% of gross annual revenue; Year 1 = $32,000 (1% of $3.2M). Amount COMPLIANT. STRUCTURAL MISCHARACTERIZATION: Code §9-5-15 requires the contribution payable to the 'City of Evanston Cannabis Community Benefit Fund' (administered by the Dept. of Community Development). The business plan frames Greenleaf as 'directing' contributions to chosen organizations and 'working with the City to identify… beneficiary organizations' — this misstates the structure (the City administers the fund; the applicant does not select beneficiaries). Code §9-5-15(D) requires an audited gross-revenue statement by an IL-licensed CPA at payment — not addressed.",
  "Compliant (amount) — discrepancy", "Compliant",
  "Clarify that the contribution is payable to the City's Cannabis Community Benefit Fund (not directly to applicant-selected organizations). Address the audited-revenue-statement (CPA) requirement for ongoing compliance."),

 ("A14", "Diversity & inclusion plan",
  "Code §9-5-5(B)(13), §9-5-20; Guide Item 14; Form §11 Q31/Q32",
  "Business plan §7.5 (vendor diversity) and §7.2 (R3 hiring) address portions. Code §9-5-20(B) requires five elements: (1) hiring practices/outreach/partnerships; (2) specific measurable goals (R3 residents, expunged/sealed-conviction individuals, other underrepresented); (3) vendor/contractor diversity; (4) training (unconscious bias, cultural competency, anti-discrimination); (5) measurement/reporting metrics. The package does NOT explicitly address unconscious-bias training, cultural competency, anti-discrimination policies, goals for expunged-conviction individuals, or measurement metrics. The standalone D&I plan (Form §11 Q32, 3–5 pp.) is NOT in the package. For social-equity applicants (Code §9-5-20(C)): the R3 hiring pipeline is described but 'specific partnerships… have not yet been formalized.'",
  "Partial — Gap", "Partial",
  "Prepare and submit the standalone D&I plan addressing ALL Code §9-5-20(B) elements, including unconscious-bias/cultural-competency training, measurable goals for expunged-conviction individuals, and reporting metrics. Formalize the workforce-development partnerships for the R3 hiring pipeline."),

 ("A15", "Environmental sustainability plan",
  "Code §9-5-5(B)(11); Guide Item 15; Form §12 Q33",
  "Business plan §5.4 covers odor and waste and references a standalone environmental plan as a separate exhibit (Form §12 Q33). The standalone Environmental Sustainability Plan is NOT in the package. Code requires energy efficiency, waste reduction, water conservation, and cannabis waste disposal (IEPA).",
  "Partial — Gap", "Partial",
  "Prepare and submit the standalone Environmental Sustainability Plan addressing energy efficiency, waste reduction/recycling, water conservation, and cannabis waste disposal consistent with IEPA regulations."),

 ("A16", "Insurance documentation",
  "Code §9-5-16; Guide Item 16; Form §8 Q27",
  "Insurance memo provided. In force: Commercial General Liability ($2M/$4M, eff. 1/15/2024) and Product Liability (bundled, $2M/$4M). Quoted/pending: Workers' Compensation (bind upon first hire). Under negotiation: Cannabis-Specific Crime/Theft ($500K, Canopy Risk — not bound; pending the final security plan). CRITICAL DISCREPANCY — Employer's Liability: Code §9-5-16(A)(3) requires Workers' Comp at statutory limits AND Employer's Liability ≥$1,000,000 per accident; the Guide (Item 16), Form (Q27), business plan (§6.4), and insurance memo (§2) ALL omit it; the lease summary §7(c) mentions employer's liability but at only $500,000/accident (below the Code's $1M). ADDITIONAL INSURED: Code §9-5-16(B) requires the City as additional insured on CGL & Product Liability — memo says GL 'can be endorsed… upon request' (not yet done). INSURER RATING (AM Best ≥A-, §9-5-16(C)) and 30-day cancellation notice (§9-5-16(D)) not addressed. No binders/commitment letters for Workers' Comp or Crime/Theft in package.",
  "Partial — Gap + discrepancy", "Partial",
  "Obtain Employer's Liability coverage ≥$1M/accident (a Code requirement omitted from the Guide/Form). Bind or obtain commitment letters for Workers' Comp and Cannabis Crime/Theft. Endorse the City as additional insured on CGL & Product Liability. Confirm insurer AM Best ratings ≥A- and 30-day cancellation-notice provisions. (The Guide/Form omission of Employer's Liability is a City-document error the applicant must still satisfy under the Code.)"),

 ("A17", "Financial viability & bonding capacity",
  "Code §9-5-5(B)(16), §9-5-17; Guide Item 17; Form §8 Q24/Q26",
  "Capitalization $1,850,000 (all equity, documented); sources/uses provided (BP §6.1/§6.2). CRITICAL GAP — Bonding capacity: Code §9-5-5(B)(16), Guide Item 17, and Form §8 Q26 all require evidence of bonding capacity (a commitment/preliminary-approval letter from a surety or financial institution) AT THE TIME OF APPLICATION. The insurance memo §5 states 'no formal commitment letter or bonding capacity letter has been obtained from Thornfield Surety Group to date' — only a 'preliminary inquiry' and 'verbal' indication. SCORING IMPACT: the Form rubric awards 0 of 10 Financial Viability points to applicants without bonding-capacity evidence.",
  "Non-Compliant — Critical gap", "Critical",
  "Obtain a formal bonding commitment letter from Thornfield Surety Group (or an irrevocable LOC commitment from Ridgeline Community Bank) BEFORE submission. Without it the applicant fails an application requirement and forfeits 10 Financial Viability points."),

 ("A18", "Social equity documentation",
  "Code §9-5-19; Guide Item 18; Form §6 Q19–Q21",
  "Okonkwo qualifies on two bases: R3 residency (tract 8092.02, on the Code §9-5-19(G) designated list) and a prior expunged cannabis conviction (2016 misdemeanor, expunged 2020). 60% ownership = majority (satisfies the Code). State license confirms the social-equity designation. DISCREPANCIES: (a) ownership threshold — Code §9-5-19(B) says 'majority' (>50%); Guide §2 and the State license say 'at least 51%.' Okonkwo's 60% satisfies both. (b) qualifying criteria — Code lists R3 residency (24-month continuous), expunged/sealed/pardoned conviction, and a family-member criterion; the Guide omits the family-member criterion and the 24-month duration; the Form adds 'Disproportionately Impacted Area' (a State, not City, criterion) and omits the family-member criterion. GAPS: 24-month continuous R3 residency not documented (BP says 'current resident' with no duration; the State license uses a different 5-of-10-years DIA standard). R3 residency proof (2 forms within 90 days + census-tract verification), the expungement court order, the signed 2-year mentorship commitment, and the signed 50% R3 hiring commitment are not in the package (Form Q21 signature block present but unexecuted).",
  "Partial — Gap + discrepancy", "Partial",
  "Document 24-month continuous residency in R3 tract 8092.02 (2 forms of proof within 90 days + census-tract verification). Provide the expungement court order/records. Obtain the signed 2-year mentorship commitment and the signed 50% R3 hiring commitment. Reconcile the social-equity criteria discrepancies (Code vs. Guide vs. Form vs. State)."),

 ("A19", "Agent-in-Charge designation",
  "Code §9-5-13(I), §9-5-11(G); Guide Item 19; Form §2 Q13",
  "Okonkwo designated Agent-in-Charge (business plan, State license). DISCREPANCY — availability standard: Code §9-5-13(I) requires the Agent-in-Charge present during all operating hours OR 'available to return to the premises within thirty (30) minutes'; Form §2 Q13(b) requires 'present… or available by telephone' and adds a 'reside within thirty (30) miles' requirement (not in the Code). The Form's 'available by telephone' is weaker than the Code's 'return within 30 minutes.' AGENT ID CARD: Code §9-5-13(E) and §9-5-11(G) require the Agent-in-Charge (and all employees) to hold a valid IL Cannabis Dispensing Organization Agent ID Card; the business plan does not confirm Okonkwo holds one (she is a licensed pharmacist, which is distinct). Form Q13 asks for the card number or 'Application Pending.'",
  "Partial — Discrepancy + gap", "Partial",
  "Confirm Okonkwo holds (or has applied for) an IL Cannabis Dispensing Organization Agent ID Card and provide the number/status. Reconcile the availability standard (Code's 30-minute return vs. Form's telephone). Okonkwo is an Evanston resident, so the Form's 30-mile requirement is satisfied."),

 ("A20", "Zoning confirmation / special use",
  "Code §9-5-9(A),(D); Guide Item 20; Form §3 Q3",
  "Location is C1a (special use required); business plan confirms. CRITICAL DISCREPANCY — permitted districts: Code §9-5-9(A) permits dispensaries (with special use) in C1a, C2, and oC1. Form §3 Q3 note lists 'C1a Commercial Mixed-Use, C1 Commercial, and B2 Business' — C1 and B2 are NOT in the Code, and C2/oC1 (which ARE in the Code) are omitted. Guide Item 20 mentions only C1a. SPECIAL-USE APPLICATION: Form §3 Q3 asks whether the applicant has applied for/obtained special-use approval; the business-plan timeline shows a ZBA hearing as part of the process but does not confirm a special-use application has been FILED. No evidence of filing.",
  "Partial — Discrepancy + gap", "Discrepancy",
  "File the special-use permit application with the ZBA (concurrent with the permit application, per the Form). Document the filing status/ZBA case number. Flag the Form's incorrect zoning-district list to the Coordinator (correct list per Code: C1a, C2, oC1)."),

 # ---- Group B ----
 ("B21", "Operating hours",
  "Code §9-5-13(F); Form §5 Q11",
  "Business plan §5.1: 6:00 AM–10:00 PM daily — COMPLIANT. DISCREPANCY: Form §5 Q11 cites 'Section 9-5-17' as the authority for hours, but operating hours are in §9-5-13(F); §9-5-17 governs the performance bond. The Form cites the wrong section.",
  "Compliant (hours) — discrepancy", "Compliant",
  "None for the hours. Flag the Form's incorrect section citation (should be §9-5-13(F)) to the Coordinator."),

 ("B22", "Maximum customer occupancy",
  "Code §9-5-13(J); Form §5 Q12",
  "CRITICAL DISCREPANCY. Code §9-5-13(J): 1 customer per 40 sq ft of RETAIL FLOOR AREA (1,600 ÷ 40 = 40 customers). Form §5 Q12: 'Total Dispensary Square Footage ÷ 40' using total premises sq ft (3,200 ÷ 40 = 80 customers). The Form's formula is WRONG — it uses total premises area instead of retail floor area, yielding double the Code's limit. The business plan §5.1 correctly uses retail floor area (1,600 ÷ 40 = 40) — COMPLIANT. If the applicant follows the Form's formula, the calculated limit (80) would violate the Code.",
  "Compliant (calc) — critical discrepancy", "Discrepancy",
  "Use retail floor area (1,600 sq ft → 40 customers) per the Code. Flag the Form's formula error to the Coordinator. Post the occupancy limit conspicuously (Code §9-5-13(J))."),

 ("B23", "Video surveillance",
  "Code §9-5-13(A); Form §5 Q15",
  "24/7, 1080p minimum, 30-day retention — business plan §5.3. COMPLIANT in concept. Code requires camera coverage of all entrances/exits, retail floor, POS, secure storage, product-handling areas, exterior (including parking/loading), and adjacent alleyways. The lease notes a rear service alley — confirm alleyway coverage in the final security plan. Recordings must be available to EPD/Coordinator within 24 hours of request.",
  "Compliant (concept) — verify", "Compliant",
  "Confirm the final security plan includes all required coverage areas (especially exterior/parking/loading and the adjacent alleyway) and the 24-hour production-on-request protocol."),

 ("B24", "Seed-to-sale / COMPASS integration",
  "Code §9-5-13(D); Form §5 Q14",
  "Business plan §5.2 confirms COMPASS integration (designed by Pinnacle). Form §5 Q14 requires a vendor confirmation letter of COMPASS integration — NOT in the package. INCONSISTENCY: lease summary §8 references 'BioTrack or successor platform' while the business plan references COMPASS; confirm the system integrates with COMPASS (the State system).",
  "Partial — Gap", "Partial",
  "Provide the tracking-system vendor's confirmation letter of COMPASS integration. Reconcile the BioTrack/COMPASS terminology."),

 ("B25", "Waste disposal plan",
  "Code §9-5-13(H); Form §5 Q17",
  "Business plan §5.4 describes waste disposal (render unusable/unrecognizable, log, licensed vendor, COMPASS recording). Referenced as a separate exhibit but the standalone plan is not fully in the package. Code requires 5-year retention of waste records.",
  "Partial", "Partial",
  "Submit the standalone waste disposal plan; confirm 5-year record retention and IEPA compliance."),

 ("B26", "Background checks",
  "Code §9-5-22; Guide Item 5; Form §10 Q14A",
  "DISCREPANCY — disqualifying-offense lookback: Code §9-5-22(C)(1) = any felony within 10 years; State license §5(d) = any other felony within 5 years. The City (10 years) is more restrictive than the State (5 years). Background-check authorizations and fingerprint receipts are NOT in the package. The State license confirms both officers passed State background checks, but the Code requires City-conducted fingerprint-based checks (IL State Police + FBI). Okonkwo's expunged misdemeanor is not disqualifying; Ramachandran has no record.",
  "Partial — Gap + discrepancy", "Partial",
  "Complete City background-check authorizations and fingerprinting for both principal officers. Note the City's 10-year felony lookback is more restrictive than the State's 5-year standard."),

 ("B27", "Permit cap / number of permits",
  "Code §9-5-4(B)",
  "Code authorizes a maximum of 4 dispensary permits. One is currently issued (Prairie Wellness). Greenleaf would be the second. COMPLIANT.",
  "Compliant", "Compliant",
  "None."),

 ("B28", "Application window",
  "Code §9-5-7(A); Guide §1; Form",
  "Window: 2/1/2024–5/31/2024. Greenleaf targets submission 5/15/2024 — within the window but with only ~16 days of buffer before it closes.",
  "Compliant (planned)", "Compliant",
  "Ensure submission on or before 5/31/2024. Given the outstanding gaps (PE cert, bonding letter, security plan, additional community session), the 5/15 target is at risk; assess whether the gaps can be cured by 5/15 or whether a later-in-window submission is needed."),

 ("B29", "Submission format",
  "Code §9-5-5(C); Guide §1; Form §2",
  "3 hard copies (bound/collated) + 1 electronic (USB, PDF). Consistent across documents. COMPLIANT in concept.",
  "Compliant (concept)", "Compliant",
  "Assemble 3 collated/bound hard copies + 1 USB PDF, tabbed by section/exhibit, delivered to the Cannabis Licensing Division."),

 ("B30", "Completeness cure period",
  "Code §9-5-7(C); Guide §1/§4; Form §13/§15",
  "CRITICAL DISCREPANCY. Code §9-5-7(C): the applicant has 30 calendar days to cure deficiencies. Guide §1/§4 and Form §13/§15: 14 calendar days. The Guide/Form are more restrictive than the Code (14 vs. 30 days), materially shortening the cure window.",
  "Discrepancy", "Discrepancy",
  "Clarify the cure period with the Coordinator. Plan for the shorter 14-day cure (more restrictive) to be safe. Flag the Code/Guide/Form conflict."),

 ("B31", "Review timeline",
  "Code §9-5-7(C); Guide §1/§4/§5; Form §6",
  "DISCREPANCY/CONFUSION. Code §9-5-7(C): completeness determination within 90 calendar days of receipt. Guide §4: completeness review targets 30 days; Guide §1/Form §6: a 90-day review period measured from the completeness determination (i.e., substantive review). The Code's 90 days = completeness determination; the Guide/Form's 90 days = substantive review after completeness. The documents use '90 days' for different milestones.",
  "Discrepancy", "Discrepancy",
  "Clarify the review timeline/milestones with the Coordinator. The Code's 90-day completeness-determination standard is controlling."),

 ("B32", "Pre-issuance conditions (overview)",
  "Code §9-5-11",
  "Post-approval conditions include: $15,000 annual permit fee; $50,000 bond/LOC; final insurance certificates (City as additional insured); Certificate of Occupancy; security-system inspection/verification; Agent ID Cards for all employees; signed community-benefit agreement; valid State license; special-use approval. CRITICAL: the valid-State-license condition (§9-5-11(I)) is at risk given the 9/21/2024 conditional-license expiration vs. anticipated Oct/Nov 2024 permit issuance (see A3).",
  "Future — critical risk flagged", "Critical",
  "Plan for all pre-issuance conditions; prioritize resolving the State-license expiration risk (file the extension petition) so a valid State license exists at local permit issuance."),

 ("B33", "Annual permit fee",
  "Code §9-5-14; Guide §6; Form",
  "$15,000, payable by certified check or EFT (Code) / cashier's check or money order (Guide) — the same payment-method discrepancy as the application fee (A2). Pre-issuance and at each annual renewal.",
  "Compliant (amount) — discrepancy", "Compliant",
  "Pay by a Code-authorized method (certified check or EFT). Reconcile the payment-method discrepancy."),

 ("B34", "Performance bond / LOC (pre-issuance)",
  "Code §9-5-17; Guide §6; Form §8 Q26",
  "$50,000; surety ≥A- AM Best or a federally insured FI; form acceptable to the City Attorney; remains 1 year after permit end; renew ≥30 days before expiration. Pre-issuance. The bonding-CAPACITY evidence (application-time) is a separate critical gap (see A17).",
  "Future — capacity gap at A17", "Future",
  "Obtain the $50,000 bond/LOC in a City-Attorney-acceptable form before issuance; ensure the surety rating ≥A- and the 1-year tail."),

 ("B35", "Records & reporting (ongoing)",
  "Code §9-5-18, §9-5-13(K)",
  "5-year record retention; quarterly self-audits (due 15 days after quarter-end); annual audited financials (CPA, within 90 days of permit-year end); annual workforce-demographics report; immediate incident reports (within 24 hours). The business plan §9.2 acknowledges the quarterly self-audits. Ongoing obligations.",
  "Future — acknowledged", "Future",
  "Establish compliance systems for all reporting; engage an IL-licensed CPA for the audited financials and the audited revenue statement."),

 ("B36", "Inspections",
  "Code §9-5-23",
  "A pre-opening inspection is required before permit issuance; ≥2 unannounced compliance inspections/year; right of inspection during operating hours without notice. Ongoing.",
  "Future", "Future",
  "Plan for the pre-opening inspection (security, building, fire, ADA); maintain an audit-ready compliance file."),

 ("B37", "Transfer restrictions",
  "Code §9-5-24",
  "No transfer of the permit; any 10%+ ownership transfer requires prior City Council approval ($2,500 fee). The lease §9 change-of-control provision (>50%) is consistent in concept but narrower than the Code's 10% threshold. Note for ongoing compliance.",
  "Future — note", "Future",
  "Be aware the Code's 10% ownership-transfer approval threshold is stricter than the lease's 50% change-of-control trigger; any future ownership change ≥10% requires City Council approval."),

 ("B38", "Prohibited acts & ongoing compliance",
  "Code §9-5-21, §9-5-25",
  "Prohibited acts (sales to minors, off-hours operation, on-premises consumption, untracked product, unlicensed agents, etc.) and penalties ($500–$5,000/day, suspension, revocation). Ongoing.",
  "Future — note", "Future",
  "Incorporate the prohibited-acts and penalty framework into staff training and the compliance program."),
]

# build matrix table
COLS = ["ID", "Compliance Requirement", "Authority (Code / Guide / Form)",
        "Findings — Applicant Submission & Assessment", "Status", "Remediation Notes"]
WIDTHS = [Inches(0.35), Inches(1.45), Inches(1.25), Inches(3.45), Inches(1.05), Inches(2.25)]

mtab = doc.add_table(rows=1, cols=6)
mtab.style = 'Table Grid'
mtab.alignment = WD_TABLE_ALIGNMENT.CENTER
mtab.autofit = False
# allow column widths
mtab_pr = mtab._tbl.tblPr
layout = OxmlElement('w:tblLayout')
layout.set(qn('w:type'), 'fixed')
mtab_pr.append(layout)

hrow = mtab.rows[0]
set_repeat_header(hrow)
for j, h in enumerate(COLS):
    c = hrow.cells[j]
    cell_text(c, h, bold=True, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade(c, '1F3A5F')
    set_cell_margins(c, top=50, bottom=50)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

current_group = None
for (rid, req, auth, find, status_label, skey, rem) in MATRIX:
    grp = "A" if rid.startswith("A") else "B"
    if grp != current_group:
        current_group = grp
        # group banner row
        grow = mtab.add_row()
        gcell = grow.cells[0]
        # merge across all 6
        for k in range(1, 6):
            gcell = gcell.merge(grow.cells[k])
        label = ("GROUP A — APPLICATION PACKAGE COMPONENTS (Submission-time requirements: "
                 "Code §9-5-5 / Guide Items 1–20)") if grp == "A" else (
                 "GROUP B — OPERATIONAL, SCORING & PROCESS REQUIREMENTS")
        cell_text(gcell, label, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(gcell, '4472C4')
        set_cell_margins(gcell, top=40, bottom=40)
    row = mtab.add_row()
    cells = row.cells
    cell_text(cells[0], rid, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(cells[1], req, bold=True, size=8)
    cell_text(cells[2], auth, size=7.5, italic=True)
    cell_text(cells[3], find, size=8)
    cell_text(cells[4], status_label, bold=True, size=8, color=STATUS_TEXT[skey],
              align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(cells[4], STATUS_COLORS[skey])
    cell_text(cells[5], rem, size=8)
    for c in cells:
        set_cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# set column widths on every cell (python-docx quirk)
for row in mtab.rows:
    for j, c in enumerate(row.cells):
        try:
            c.width = WIDTHS[j]
        except Exception:
            pass

# ============================================================
# 5. back to portrait — discrepancy log + priority actions
# ============================================================
port = doc.add_section(WD_SECTION.NEW_PAGE)
port.orientation = WD_ORIENT.PORTRAIT
port.page_width = Inches(8.5)
port.page_height = Inches(11)
port.top_margin = Inches(0.7)
port.bottom_margin = Inches(0.7)
port.left_margin = Inches(0.8)
port.right_margin = Inches(0.8)

add_heading(doc, '5.  Cross-Document Discrepancy Log (Code vs. Guide vs. Form)', level=1)
body(doc,
     'The following conflicts exist among the City\'s own governing documents. Because the Coordinator '
     'administers the process using the Guide and Form, each conflict creates ambiguity the applicant must '
     'resolve (in writing, ideally) before submission. The Code is legally controlling; where the Guide or '
     'Form is more restrictive, satisfying the stricter standard is the safer course.', size=9.5, space_after=6)

DISC = [
 ("D1", "Completeness cure period", "Code §9-5-7(C): 30 calendar days", "Guide §1/§4 & Form §13/§15: 14 calendar days",
  "Plan for a 14-day cure (stricter). Confirm in writing with the Coordinator; the Code entitles 30 days."),
 ("D2", "Application-fee payment method", "Code §9-5-6(C): certified check or EFT", "Guide: cashier's check or money order; Form: certified check, cashier's check, or wire transfer",
  "Pay by certified check or EFT (Code-authoritative). Money order (Guide) is not Code-authorized."),
 ("D3", "Completeness/review timeline", "Code §9-5-7(C): completeness determination within 90 days of receipt", "Guide §4: 30-day completeness target; Guide §1/Form §6: 90-day review from completeness determination",
  "The Code's 90-day completeness standard controls. Clarify milestones with the Coordinator."),
 ("D4", "Buffer-distance measurement", "Code §9-5-9(C): nearest property line to nearest property line (straight-line)", "Form §3 Q7: nearest pedestrian route; straight-line 'not acceptable'",
  "Obtain pedestrian-route measurements (or confirm straight-line acceptable). Distances likely still compliant given margins."),
 ("D5", "Permitted zoning districts", "Code §9-5-9(A): C1a, C2, oC1 (special use)", "Form §3 Q3: C1a, C1, B2; Guide Item 20: C1a only",
  "The Form lists C1/B2 (not in the Code) and omits C2/oC1. Use the Code list; flag the Form error to the Coordinator."),
 ("D6", "Maximum-occupancy formula", "Code §9-5-13(J): retail floor area ÷ 40 (= 40)", "Form §5 Q12: total premises sq ft ÷ 40 (= 80)",
  "Use retail floor area per the Code (40 customers). The Form's formula doubles the lawful limit."),
 ("D7", "Employer's Liability insurance", "Code §9-5-16(A)(3): Workers' Comp statutory + Employer's Liability ≥$1M/accident", "Guide Item 16 & Form Q27: Workers' Comp statutory only (Employer's Liability omitted)",
  "Obtain Employer's Liability ≥$1M/accident. The Guide/Form omission is a City-document error the applicant must still satisfy."),
 ("D8", "Social-equity ownership threshold", "Code §9-5-19(B): 'majority' of ownership (>50%)", "Guide §2 & State license: 'at least 51%'",
  "Okonkwo's 60% satisfies both. Note the differing articulation; the Code's 'majority' controls."),
 ("D9", "Social-equity qualifying criteria", "Code §9-5-19(B): R3 residency (24-mo continuous); expunged/sealed/pardoned conviction; family-member criterion", "Guide: R3 residency + arrest/conviction/adjudication (omits family member & 24-mo); Form: adds 'Disproportionately Impacted Area' (a State criterion), omits family member; State: 5-of-10-yr DIA + arrest/conviction/adjudication + family member",
  "Document eligibility under the Code criteria (24-mo R3 residency; expungement records). Do not rely solely on the State DIA standard."),
 ("D10", "Certificate of Good Standing recency", "Code §9-5-5(B)(3) & Guide Item 4: within 60 days", "Form §13 checklist: within 30 days",
  "Obtain within 30 days to satisfy both; the Form is stricter."),
 ("D11", "Lease-term measurement basis", "Code §9-5-5(B)(5): remaining term from anticipated permit issuance", "Guide Item 6 & Form §3 Q4: from application submission",
  "No practical impact (7-year lease). Note the differing basis."),
 ("D12", "Agent-in-Charge availability", "Code §9-5-13(I): present or 'available to return within 30 minutes'", "Form §2 Q13(b): 'present or available by telephone'; adds 30-mile residency",
  "Meet the Code's 30-minute-return standard. The Form's telephone standard is weaker; the 30-mile residency is a Form-only addition (Okonkwo satisfies it)."),
 ("D13", "Operating-hours citation", "Code §9-5-13(F) (hours)", "Form §5 Q11 cites '§9-5-17' (the performance-bond section)",
  "Cite §9-5-13(F). Flag the Form's wrong citation to the Coordinator."),
 ("D14", "SSN disclosure", "Code §9-5-5(B)(4): full SSN for each principal officer", "Form Q14A: last 4 digits only",
  "Confirm with the Coordinator whether full SSN is required; the Code text requires it."),
 ("D15", "Felony lookback (background checks)", "Code §9-5-22(C)(1): any felony within 10 years", "State license §5(d): any other felony within 5 years",
  "The City (10 years) is stricter than the State (5 years). Screen principal officers against the 10-year standard."),
 ("D16", "Coordinator contact information", "Guide: (847) 448-4311 / rvasquezmorris@cityofevanston.org", "Form: (847) 448-8163 / rvasquez-morris@cityofevanston.org (hyphen)",
  "Confirm the correct phone/email with the City; the documents disagree."),
]

dtab = doc.add_table(rows=1, cols=5)
dtab.style = 'Table Grid'
dtab.alignment = WD_TABLE_ALIGNMENT.CENTER
dh = dtab.rows[0].cells
for j, h in enumerate(["#", "Topic", "Code (controlling)", "Guide / Form", "Recommended resolution"]):
    cell_text(dh[j], h, bold=True, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade(dh[j], '1F3A5F')
    set_cell_margins(dh[j])
DW = [Inches(0.3), Inches(1.25), Inches(1.7), Inches(1.7), Inches(2.55)]
for (num, topic, codev, gf, res) in DISC:
    row = dtab.add_row().cells
    cell_text(row[0], num, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row[1], topic, bold=True, size=8)
    cell_text(row[2], codev, size=8)
    cell_text(row[3], gf, size=8)
    cell_text(row[4], res, size=8)
    for c in row:
        set_cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
for row in dtab.rows:
    for j, c in enumerate(row.cells):
        try:
            c.width = DW[j]
        except Exception:
            pass

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# applicant-internal inconsistencies
add_heading(doc, '6.  Applicant-Internal Inconsistencies to Reconcile', level=1)
body(doc, 'The following inconsistencies appear within the applicant\'s own package and should be corrected '
     'before submission to avoid a completeness challenge or a material-misrepresentation concern:', size=9.5, space_after=4)
bullet(doc, 'business plan & lease summary say "Northwest Statutory Services"; the State license says "Keystone Statutory Services" (same address). Confirm the correct registered-agent name and correct all documents.', bold_lead='Registered-agent name — ')
bullet(doc, 'business plan lists (872) 555-0193; the insurance memo lists (847) 555-0193. Use one consistent contact number.', bold_lead='Contact telephone number — ')
bullet(doc, 'lease summary §8 references "BioTrack or successor platform"; the business plan references COMPASS. Confirm the system integrates with the State COMPASS platform and use consistent terminology.', bold_lead='Seed-to-sale system — ')
bullet(doc, 'the business plan states the dispensary will serve both adult-use and medical customers, but the community-engagement notices describe only an "adult-use cannabis dispensary." Align the notices/plan language.', bold_lead='Adult-use vs. medical scope — ')

# critical priority actions
add_heading(doc, '7.  Critical Priority Remediation Actions', level=1)
body(doc, 'These items must be resolved before submission; each, if uncured, renders the application '
     'incomplete or causes material scoring loss. They are listed in suggested priority order.', size=9.5, space_after=4)
PRI = [
 ("1", "Obtain the odor-mitigation PE certification (A10).", "Finalize the HVAC/carbon-filtration design and secure the signed/sealed IL-PE certification. The Form deems the application incomplete without it; it cannot be deferred to pre-issuance."),
 ("2", "Hold an additional compliant community-engagement session (A12).", "Session 2 was noticed 12 days in advance (< 14). Per Code §9-5-8(F), hold an additional session with full ≥14-day newspaper + City-website notice; add a telephone number to the notice; obtain publisher's affidavits and website-posting confirmations."),
 ("3", "Obtain a bonding-capacity commitment letter (A17).", "Secure a formal commitment letter from Thornfield Surety Group (or an LOC commitment from Ridgeline Community Bank). Without it the applicant forfeits 10 Financial Viability points and fails an application requirement."),
 ("4", "Resolve the State-license expiration risk (A3 / B32).", "The conditional license expires 9/21/2024, before the anticipated local permit issuance. File a State extension petition by ~7/23/2024 (≥60 days before expiration) and expedite the local process so a valid State license exists at local permit issuance (Code §9-5-11(I))."),
 ("5", "Finalize the security plan (A9).", "Complete the Pinnacle security plan (it also unblocks the cannabis crime/theft underwriting). Ensure all Code §9-5-13 coverage areas are addressed, including the adjacent rear service alley."),
 ("6", "Assemble the missing required exhibits (A4, A5, A7, A8, A14, A15, A18, A19, A20, A24).", "Certificate of Good Standing; IRS EIN letter; Articles & Operating Agreement; individual disclosure forms/background-check authorizations; scaled architect/engineer site plan; updated survey; standalone D&I, environmental, and waste-disposal plans; COMPASS vendor letter; social-equity residency/expungement/commitment documents; Agent ID Card confirmation; special-use application filing."),
 ("7", "Reconcile the insurance posture (A16).", "Add Employer's Liability ≥$1M/accident (Code requirement omitted from the Guide/Form); bind or commit Workers' Comp and Cannabis Crime/Theft; endorse the City as additional insured; confirm AM Best ≥A- ratings and 30-day cancellation notice."),
 ("8", "Resolve the buffer-measurement and occupancy-formula conflicts (A8, B22).", "Obtain pedestrian-route buffer measurements (or confirm straight-line acceptable) and an updated survey; use retail floor area (not total premises) for the occupancy calculation per the Code."),
]
ptab = doc.add_table(rows=1, cols=3)
ptab.style = 'Table Grid'
ptab.alignment = WD_TABLE_ALIGNMENT.CENTER
ph = ptab.rows[0].cells
for j, h in enumerate(["#", "Action", "Why it is critical"]):
    cell_text(ph[j], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade(ph[j], '9C0006')
    set_cell_margins(ph[j])
PW = [Inches(0.3), Inches(2.6), Inches(4.6)]
for (num, act, why) in PRI:
    row = ptab.add_row().cells
    cell_text(row[0], num, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x9C,0x00,0x06))
    cell_text(row[1], act, bold=True, size=9)
    cell_text(row[2], why, size=9)
    for c in row:
        set_cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
for row in ptab.rows:
    for j, c in enumerate(row.cells):
        try:
            c.width = PW[j]
        except Exception:
            pass

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# closing
add_heading(doc, '8.  Reviewer\'s Note & Limitations', level=1)
body(doc,
     'This matrix is based on the eight documents listed in Section 3 as provided. Several items the applicant '
     'represents as "to be included in the full application package" (e.g., exhibits A–E to the lease, the '
     'organizational chart, the standalone security/odor/environmental/D&I/waste plans, the COMPASS vendor '
     'letter, and the social-equity supporting documents) were not present in the package reviewed and are '
     'therefore marked as gaps; their status should be re-verified once the complete package is assembled. '
     'Where the Code, Guide, and Form conflict, the Code is the controlling legal authority, but the applicant '
     'should obtain written confirmation from the Cannabis Licensing Coordinator on each conflict before '
     'reliance, and should satisfy the stricter standard where feasible. This review is a compliance-readiness '
     'assessment and does not constitute a City completeness determination or legal advice.', size=9.5, space_after=6)

body(doc, 'End of Compliance Requirements Matrix.', italic=True, size=9, space_after=2)

out = 'requirements-matrix.docx'
doc.save(out)
print('Saved', out)