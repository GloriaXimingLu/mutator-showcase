#!/usr/bin/env python3
"""Generate motion-to-dismiss.docx, proposed-order.docx, cover-memo.docx
for Arcadia Health Systems, LLC v. Meridian Cloud Solutions, Inc.,
No. 1:23-cv-00847-CMA (W.D. Tex. Austin Division).

Formatting per Judge Alvarez's Standing Order:
  - 12pt Times New Roman; double-spaced body; 1" margins; page numbers bottom-center.
"""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
BODY_SIZE = 12
FOOTNOTE_SIZE = 10

# ----------------------------------------------------------------------------
# Low-level helpers
# ----------------------------------------------------------------------------

def _set_run_font(run, size=BODY_SIZE, bold=False, italic=False, underline=False, font=FONT):
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def _para_format(p, align=None, space_before=0, space_after=0, line=2.0, left_indent=None,
                 right_indent=None, first_line=None, keep_with_next=False):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if right_indent is not None:
        pf.right_indent = Inches(right_indent)
    if first_line is not None:
        pf.first_line_indent = Inches(first_line)
    pf.keep_with_next = keep_with_next


def add_rich(p, text, size=BODY_SIZE, base_bold=False, base_italic=False):
    """Add runs to paragraph p, parsing **bold** and *italic* markers."""
    # Convert dash shorthand: --- -> em dash, -- -> en dash (page ranges)
    text = text.replace('---', '\u2014').replace('--', '\u2013')
    # Tokenize on ** and *
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = p.add_run(tok[2:-1])
            _set_run_font(run, size=size, bold=True, italic=base_italic)
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            run = p.add_run(tok[1:-1])
            _set_run_font(run, size=size, bold=base_bold, italic=True)
        else:
            run = p.add_run(tok)
            _set_run_font(run, size=size, bold=base_bold, italic=base_italic)


def body_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=0.5, space_after=0):
    p = doc.add_paragraph()
    _para_format(p, align=align, line=2.0, first_line=first_line, space_after=space_after)
    add_rich(p, text)
    return p


def flush_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0, bold=False):
    p = doc.add_paragraph()
    _para_format(p, align=align, line=2.0, first_line=None, space_after=space_after)
    add_rich(p, text, base_bold=bold)
    return p


def block_quote(doc, text):
    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.0,
                 left_indent=0.5, right_indent=0.5, space_before=6, space_after=6)
    add_rich(p, text)
    return p


def heading(doc, text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    if level == 1:
        _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=2.0,
                     space_before=space_before, space_after=space_after, keep_with_next=True)
        add_rich(p, text, base_bold=True)
    else:
        _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=2.0,
                     space_before=space_before, space_after=space_after, keep_with_next=True)
        add_rich(p, text, base_bold=True, base_italic=(level == 3))
    return p


def center_line(doc, text, bold=False, size=BODY_SIZE, space_after=0, line=1.15):
    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=line, space_after=space_after)
    add_rich(p, text, size=size, base_bold=bold)
    return p


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, space_after=0)
    run = p.add_run()
    _set_run_font(run, size=BODY_SIZE)
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def setup_document(doc):
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1); s.right_margin = Inches(1)
        add_page_number_footer(s)
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'nil')
        borders.append(e)
    tblPr.append(borders)


def caption(doc, title_text, response_deadline=None):
    """Standard W.D. Tex. caption with § brackets."""
    center_line(doc, "UNITED STATES DISTRICT COURT", bold=True, space_after=0, line=1.15)
    center_line(doc, "WESTERN DISTRICT OF TEXAS", bold=True, space_after=0, line=1.15)
    center_line(doc, "AUSTIN DIVISION", bold=True, space_after=6, line=1.15)

    rows = [
        ("ARCADIA HEALTH SYSTEMS, LLC,", ""),
        ("", ""),
        ("             Plaintiff,", ""),
        ("", ""),
        ("v.", "Case No. 1:23-cv-00847-CMA"),
        ("", ""),
        ("MERIDIAN CLOUD SOLUTIONS, INC.,", ""),
        ("", ""),
        ("             Defendant.", ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _no_borders(table)
    # column widths
    for r in table.rows:
        r.cells[0].width = Inches(4.2)
        r.cells[1].width = Inches(2.3)
    for i, (left, right) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        p0 = c0.paragraphs[0]; _para_format(p0, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
        if left:
            add_rich(p0, left)
        p1 = c1.paragraphs[0]; _para_format(p1, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
        # § symbol on every row
        sec = p1.add_run("§  ")
        _set_run_font(sec, size=BODY_SIZE)
        if right:
            add_rich(p1, right, base_bold=True)
    # spacing after table
    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=2)

    # Document title
    center_line(doc, title_text, bold=True, space_after=2, line=1.15)
    if response_deadline:
        center_line(doc, response_deadline, bold=False, size=11, space_after=8, line=1.15)
    else:
        sp2 = doc.add_paragraph(); _para_format(sp2, line=1.0, space_after=4)
    # horizontal rule
    hr = doc.add_paragraph(); _para_format(hr, line=1.0, space_after=6)
    pPr = hr._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom); pPr.append(pbdr)


def signature_block(doc, signers, date_line, firm_lines, party_line):
    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=6)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
    add_rich(p, "Respectfully submitted,")
    p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
    add_rich(p, firm_lines[0], base_bold=True)
    p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
    add_rich(p, "By: /s/ " + signers[0]['name'])
    for s in signers:
        p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
        add_rich(p, f"{s['name']}")
        p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
        add_rich(p, f"{s['bar']}")
    for fl in firm_lines[1:]:
        p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
        add_rich(p, fl)
    p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
    add_rich(p, party_line, base_italic=True)
    p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
    add_rich(p, date_line)


# Counsel info (from Meridian's Answer, the operative pleading)
MERIDIAN_COUNSEL = [
    {'name': 'Margaret "Meg" Calloway', 'bar': 'Texas State Bar No. 24037891'},
    {'name': 'David Arsenault', 'bar': 'Texas State Bar No. 24098254'},
]
FIRM_LINES = [
    "STONEBRIDGE & CALLOWAY LLP",
    "600 Congress Avenue, Suite 2800",
    "Austin, Texas 78701",
    "Telephone: (512) 555-4200  Facsimile: (512) 555-4201",
    "Email: mcalloway@stonebridgecalloway.com",
    "Email: darsenault@stonebridgecalloway.com",
]
PARTY_LINE = "Attorneys for Defendant Meridian Cloud Solutions, Inc."

PLAINTIFF_COUNSEL = ("Jonathan Breckenridge", "Wexford Hale LLP",
                     "300 West 6th Street, Suite 1500, Austin, TX 78701",
                     "jbreckenridge@wexfordhale.com")


# ============================================================================
# TABLE OF AUTHORITIES (for the motion)
# ============================================================================

# Each entry: (display, search_key). search_key is used to locate page numbers
# in the rendered PDF (searched in body pages only, after the TOA).
TOA_DATA = [
    ("Cases", [
        ("*Ashcroft v. Iqbal*, 556 U.S. 662 (2009)", "556 U.S. 662"),
        ("*Bell Atl. Corp. v. Twombly*, 550 U.S. 544 (2007)", "550 U.S. 544"),
        ("*Benchmark Elecs., Inc. v. J.M. Huber Corp.*, 343 F.3d 719 (5th Cir. 2003)", "343 F.3d"),
        ("*Brasby v. Morris Dynamics, Inc.*, 947 A.2d 1042 (Del. 2008)", "947 A.2d"),
        ("*Castrol Inc. v. Pennzoil Co.*, 987 F.2d 939 (3d Cir. 1993)", "987 F.2d"),
        ("*Chapman Custom Homes, Inc. v. Dallas Plumbing Co.*, 445 S.W.3d 716 (Tex. 2014)", "445 S.W.3d"),
        ("*Dorsey v. Portfolio Equities, Inc.*, 540 F.3d 333 (5th Cir. 2008)", "540 F.3d"),
        ("*Dresser-Rand Co. v. Virtual Automation, Inc.*, 361 F.3d 831 (5th Cir. 2004)", "361 F.3d"),
        ("*Eagle Indus., Inc. v. DeVilbiss Health Care, Inc.*, 702 A.2d 1228 (Del. 1997)", "702 A.2d"),
        ("*Excess Underwriters at Lloyd's v. Frank's Casing Crew & Rental Tools, Inc.*, 246 S.W.3d 42 (Tex. 2008)", "246 S.W.3d"),
        ("*Federal Land Bank Ass'n v. Sloane*, 825 S.W.2d 439 (Tex. 1992)", "825 S.W.2d"),
        ("*Flaherty & Crumrine Preferred Income Fund, Inc. v. TXU Corp.*, 565 F.3d 200 (5th Cir. 2009)", "565 F.3d 200"),
        ("*Fortune Prod. Co. v. Conoco, Inc.*, 52 S.W.3d 671 (Tex. 2000)", "52 S.W.3d"),
        ("*Kana Software, Inc. v. Sealand Tech., Inc.*, 178 A.3d 1045 (Del. Ch. 2017)", "178 A.3d"),
        ("*Kuhn Constr., Inc. v. Diamond State Port Corp.*, 990 A.2d 393 (Del. 2010)", "990 A.2d"),
        ("*Lormand v. US Unwired, Inc.*, 565 F.3d 228 (5th Cir. 2009)", "565 F.3d 228"),
        ("*McCamish, Martin, Brown & Loeffler v. F.E. Appling Interests*, 991 S.W.2d 787 (Tex. 1999)", "991 S.W.2d"),
        ("*Melody Home Mfg. Co. v. Barnes*, 741 S.W.2d 349 (Tex. 1987)", "741 S.W.2d"),
        ("*Pizza Hut, Inc. v. Papa John's Int'l, Inc.*, 227 F.3d 489 (5th Cir. 2000)", "227 F.3d"),
        ("*PPG Indus., Inc. v. JMB/Houston Ctrs. Partners Ltd.*, 146 S.W.3d 79 (Tex. App.\u2014Houston [1st Dist.] 2004, no pet.)", "146 S.W.3d"),
        ("*Precision Healthcare Sols. v. Nextera Data Sys.*, 458 F. Supp. 3d 544 (N.D. Tex. 2020)", "458 F. Supp. 3d"),
        ("*Presidio Enters., Inc. v. Warner Bros. Distrib. Corp.*, 784 F.2d 674 (5th Cir. 1986)", "784 F.2d"),
        ("*Sharyland Water Supply Corp. v. City of Alton*, 354 S.W.3d 407 (Tex. 2011)", "354 S.W.3d"),
        ("*SIGA Techs., Inc. v. PharmAthene, Inc.*, 67 A.3d 330 (Del. 2013)", "67 A.3d"),
        ("*Simulados, Inc. v. Canton Health Mgmt. Co.*, 2019 WL 4573218 (W.D. Tex. Sept. 20, 2019) (unpublished)", "2019 WL 4573218"),
    ]),
    ("Statutes and Rules", [
        ("Fed. R. Civ. P. 9(b)", "__PASSIM__"),
        ("Fed. R. Civ. P. 12(b)(6)", "__PASSIM__"),
        ("Tex. Bus. & Com. Code \u00a7 17.49(f)", "17.49(f)"),
    ]),
    ("Other Authorities", [
        ("Restatement (Second) of Torts \u00a7 552", "Restatement (Second) of Torts"),
    ]),
]


def build_toa(doc, toa_pages):
    """Insert a Table of Authorities. toa_pages maps search_key -> page string."""
    center_line(doc, "TABLE OF AUTHORITIES", bold=True, space_after=6, line=1.15)
    for cat, entries in TOA_DATA:
        p = doc.add_paragraph()
        _para_format(p, line=1.15, space_before=6, space_after=2, keep_with_next=True)
        add_rich(p, cat, base_bold=True)
        for display, key in entries:
            pages = toa_pages.get(key, "")
            if key == "__PASSIM__":
                pages = "passim"
            p = doc.add_paragraph()
            _para_format(p, line=1.15, space_after=2, left_indent=0.25)
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            add_rich(p, display)
            r = p.add_run("\t" + pages)
            _set_run_font(r)
    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=6)
    # horizontal rule
    hr = doc.add_paragraph(); _para_format(hr, line=1.0, space_after=6)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom); pPr.append(pbdr)


def compute_toa_pages(docx_path):
    """Render docx to PDF and compute, for each TOA search_key, the body pages
    on which it appears (excluding the TOA/caption pages)."""
    import subprocess, tempfile, os
    tmp = tempfile.mkdtemp()
    subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', tmp, docx_path],
                   capture_output=True, timeout=120)
    pdf = os.path.join(tmp, os.path.basename(docx_path).replace('.docx', '.pdf'))
    res = subprocess.run(['pdftotext', '-layout', pdf, '-'], capture_output=True, text=True, timeout=60)
    pages = res.stdout.split('\f')
    # find first body page (where CERTIFICATE OF CONFERENCE heading appears)
    cert_idx = None
    for i, pg in enumerate(pages):
        if 'CERTIFICATE OF CONFERENCE' in pg:
            cert_idx = i
            break
    if cert_idx is None:
        cert_idx = 0

    def body_text(i):
        if i < cert_idx:
            return ""
        if i == cert_idx:
            idx = pages[i].find('CERTIFICATE OF CONFERENCE')
            raw = pages[i][idx:]
        else:
            raw = pages[i]
        # collapse whitespace (incl. line breaks) so citations that wrap across
        # lines in the PDF still match the search key
        return re.sub(r'\s+', ' ', raw)

    toa_pages = {}
    for cat, entries in TOA_DATA:
        for display, key in entries:
            if key == "__PASSIM__":
                continue
            found = []
            for i in range(cert_idx, len(pages)):
                if key in body_text(i):
                    found.append(str(i + 1))
            toa_pages[key] = ", ".join(found)
    return toa_pages


# ============================================================================
# DOCUMENT 1: MOTION TO DISMISS
# ============================================================================

def build_motion(toa_pages=None):
    doc = Document()
    setup_document(doc)
    caption(doc,
            "DEFENDANT MERIDIAN CLOUD SOLUTIONS, INC.'S MOTION TO DISMISS "
            "PLAINTIFF'S FIRST AMENDED COMPLAINT PURSUANT TO FEDERAL RULE OF "
            "CIVIL PROCEDURE 12(b)(6)",
            response_deadline="Pursuant to the Court's Standing Order, Plaintiff's response is due "
                              "within 21 days of the filing of this motion.")

    # Table of Authorities
    build_toa(doc, toa_pages or {})

    # Certificate of Conference (preliminary)
    heading(doc, "CERTIFICATE OF CONFERENCE", level=1, space_before=0)
    body_para(doc,
        "Pursuant to Section 3.1 of the Court's Standing Order for Civil Cases, counsel for "
        "Defendant Meridian Cloud Solutions, Inc. (\u201cMeridian\u201d) conferred in good faith with "
        "counsel for Plaintiff Arcadia Health Systems, LLC (\u201cArcadia\u201d) regarding the relief "
        "sought in this Motion. The parties are unable to agree on the dismissal of any claim. "
        "Meridian notified Arcadia's counsel of its intent to file a Rule 12(b)(6) motion by letter "
        "dated November 8, 2023, and the parties have since conferred by telephone and email. "
        "Arcadia opposes dismissal of all counts. Accordingly, further conferral would be futile, "
        "and Meridian submits this Motion for the Court's consideration.")

    # I. INTRODUCTION
    heading(doc, "I.  INTRODUCTION", level=1)
    body_para(doc,
        "This is a case about a sophisticated commercial party that struck a hard-fought bargain, "
        "received the benefit of that bargain, and now\u2014dissatisfied with the result\u2014seeks to "
        "rewind the deal through five overlapping claims for $47.3 million. Arcadia Health Systems, "
        "LLC is a $62-million-per-year healthcare IT company that was represented by experienced "
        "outside counsel throughout the negotiation of a $14.7 million Master Software License and "
        "Services Agreement (the \u201cMSLSA\u201d). Arcadia conducted its own due diligence, "
        "negotiated the contract's risk-allocation provisions, and signed a fully integrated "
        "agreement that disclaimed reliance on pre-contractual representations, capped liability, "
        "waived consequential damages, and channeled any product concerns through a limited "
        "warranty with an exclusive remedy.")
    body_para(doc,
        "The First Amended Complaint (the \u201cFAC\u201d) nevertheless asks this Court to disregard "
        "that bargain. Arcadia repackages a single grievance\u2014that the NexusCore software did not "
        "perform to Arcadia's expectations\u2014into five counts: breach of contract, fraud, negligent "
        "misrepresentation, unjust enrichment, and a claim under the Texas Deceptive Trade "
        "Practices\u2013Consumer Protection Act (the \u201cDTPA\u201d). Each count fails as a matter of "
        "law. The DTPA claim is statutorily exempt because the transaction exceeds $500,000 and "
        "Arcadia is not an individual. The unjust-enrichment claim is barred because an express "
        "contract governs the same subject matter. The fraud and negligent-misrepresentation "
        "claims fail under Rule 9(b), the economic-loss doctrine, the puffery doctrine, and the "
        "MSLSA's integration clause. And the breach-of-contract claim is defeated by the MSLSA's "
        "deemed-acceptance provision, the expired limited warranty, and the contractual limitations "
        "on remedies and damages.")
    body_para(doc,
        "Meridian therefore moves, pursuant to Federal Rule of Civil Procedure 12(b)(6), to dismiss "
        "all five counts of the FAC with prejudice.")

    # II. FACTUAL BACKGROUND
    heading(doc, "II.  FACTUAL BACKGROUND", level=1)
    body_para(doc,
        "The Court may consider the MSLSA, Statement of Work #1 (\u201cSOW-1\u201d), the four Change "
        "Orders (CO-001 through CO-004), and Meridian's December 9, 2021 written proposal on this "
        "Motion because each is referenced in and central to Arcadia's claims. *Lormand v. US "
        "Unwired, Inc.*, 565 F.3d 228, 255\u201356 (5th Cir. 2009). The FAC itself repeatedly invokes "
        "these documents and lists them in its exhibit list. FAC \u00b6\u00b6 33\u201347; FAC Ex. List "
        "(Exhibits A\u2013C, H).")

    heading(doc, "A.  The parties are sophisticated commercial entities that negotiated a "
                 "comprehensive, integrated agreement.", level=2)
    body_para(doc,
        "Arcadia is a regional healthcare IT services provider with annual revenue of approximately "
        "$62 million. FAC \u00b6 6. Meridian is a publicly traded enterprise software company with "
        "annual revenue of approximately $380 million. FAC \u00b6 9. The parties entered into the "
        "MSLSA on March 15, 2022, for a total contract value of $14,700,000. FAC \u00b6\u00b6 33\u201334. "
        "The FAC itself admits that Arcadia engaged outside counsel to negotiate the MSLSA and that "
        "Arcadia \u201cma[d]e efforts to negotiate more favorable terms, including enhanced limitation "
        "of liability provisions and expanded warranty protections.\u201d FAC \u00b6 32. The FAC further "
        "admits that Arcadia's CTO \u201cconducted certain independent due diligence, including "
        "reference calls with two existing Meridian healthcare clients.\u201d FAC \u00b6 31.")
    body_para(doc,
        "The MSLSA is a fully integrated agreement. Section 12.1 provides that the MSLSA "
        "\u201cconstitutes the entire agreement between the Parties with respect to the subject matter "
        "hereof and supersedes all prior and contemporaneous agreements, proposals, representations, "
        "and understandings, whether oral or written,\u201d and expressly supersedes all prior "
        "\u201cmarketing materials, product descriptions, brochures, presentations, and "
        "demonstrations.\u201d MSLSA \u00a7 12.1. Section 12.1 further states that \u201c[e]ach Party "
        "acknowledges that, in entering into this Agreement, it has not relied upon any statement, "
        "representation, warranty, or agreement of the other Party except as expressly set forth in "
        "this Agreement\u201d and that \u201c[e]ach Party acknowledges that it has been represented by "
        "counsel in the negotiation and execution of this Agreement.\u201d Id.")

    heading(doc, "B.  The MSLSA comprehensively allocates the risk of product nonconformance.", level=2)
    body_para(doc,
        "The MSLSA contains a detailed, negotiated risk-allocation framework. Section 9.1 provides a "
        "limited warranty that the software will \u201cperform substantially in accordance with the "
        "Documentation\u201d for ninety (90) days following Go-Live, and designates repair, "
        "replacement, or a pro-rata refund as the \u201csole and exclusive remedy\u201d for any breach. "
        "MSLSA \u00a7 9.1. Section 9.4 disclaims all other warranties and states that the software is "
        "provided \u201cAS IS\u201d and \u201cAS AVAILABLE,\u201d that Arcadia \u201chas relied on its own "
        "due diligence and evaluation in selecting the Licensed Software,\u201d and that Meridian does "
        "not warrant the software will \u201coperate without interruption or be error-free.\u201d MSLSA "
        "\u00a7 9.4. Section 8.1 caps each party's aggregate liability at the fees paid or payable in "
        "the twelve months preceding the event giving rise to liability. MSLSA \u00a7 8.1. Section 8.2 "
        "excludes all \u201cindirect, incidental, special, consequential, punitive, or exemplary "
        "damages, including \u2026 damages for loss of profits, goodwill, use, data, or other "
        "intangible losses.\u201d MSLSA \u00a7 8.2. Section 8.3 designates the MSLSA's remedies as "
        "\u201csole and exclusive.\u201d MSLSA \u00a7 8.3.")
    body_para(doc,
        "Section 5.3 establishes a deemed-acceptance regime: Arcadia had thirty (30) days following "
        "deployment to conduct acceptance testing and deliver written notice of any material "
        "nonconformity; absent such written notice, the software is \u201cdeemed accepted.\u201d MSLSA "
        "\u00a7 5.3(d). Finally, the MSLSA and SOW-1 allocate all data-migration responsibility to "
        "Arcadia. SOW-1 \u00a7 3.2; MSLSA Ex. C \u00a7 C.2. Meridian's data-migration role was limited "
        "to \u201cadvisory consulting services,\u201d and Meridian bore \u201cno responsibility for the "
        "accuracy, completeness, or integrity of migrated data.\u201d Id.")

    heading(doc, "C.  Arcadia accepted the software and its data-migration vendor caused the "
                 "problems it now complains of.", level=2)
    body_para(doc,
        "NexusCore achieved Go-Live on January 15, 2023. FAC \u00b6 45. Under Section 5.3, the "
        "thirty-day acceptance-testing window closed on February 14, 2023. Arcadia did not deliver "
        "any written notice of material nonconformity within that period. Arcadia's first written "
        "communication\u2014Dr. Okonkwo's March 8, 2023 email to Meridian's Vice President of Customer "
        "Success\u2014arrived twenty-two days after the window closed and reported only "
        "\u201cintermittent latency issues\u201d and \u201coccasional report generation errors,\u201d "
        "noting that \u201c[n]either of these issues has caused any patient-safety concerns.\u201d FAC "
        "\u00b6 50. Acceptance was therefore deemed granted as of February 14, 2023. MSLSA \u00a7 5.3(d).")
    body_para(doc,
        "The data-integrity problems Arcadia now attributes to Meridian were, by Arcadia's own "
        "admission, caused by its independently retained systems integrator, Linden Park Consulting, "
        "LLC. Change Order CO-004\u2014signed by Dr. Okonkwo on Arcadia's behalf\u2014expressly "
        "acknowledges that the data-integrity issues \u201carose from errors in data migration scripts "
        "prepared and executed by Linden Park Consulting, LLC, and not from any defect in the "
        "NexusCore software or any act or omission of [Meridian].\u201d CO-004 \u00a7 5.1. This "
        "admission directly contradicts the FAC's allegation that \u201cMeridian and its agents\u201d "
        "failed the data migration. FAC \u00b6\u00b6 43, 52. Where a complaint's allegations contradict "
        "the terms of a document referenced in and central to the complaint, the document controls. "
        "The four Change Orders added $2,350,000 in scope, bringing total consideration to "
        "$17,050,000. FAC \u00b6\u00b6 46\u201347.")

    # III. LEGAL STANDARD
    heading(doc, "III.  LEGAL STANDARD", level=1)
    body_para(doc,
        "A motion to dismiss under Rule 12(b)(6) tests whether the complaint \u201ccontain[s] "
        "sufficient factual matter, accepted as true, to state a claim to relief that is plausible "
        "on its face.\u201d *Ashcroft v. Iqbal*, 556 U.S. 662, 678 (2009) (quoting *Bell Atl. Corp. v. "
        "Twombly*, 550 U.S. 544, 570 (2007)). The Court accepts all well-pleaded factual allegations "
        "as true and draws all reasonable inferences in the plaintiff's favor, but it is not bound "
        "to accept \u201cthreadbare recitals of the elements of a cause of action, supported by mere "
        "conclusory statements.\u201d *Id.* at 678. \u201c[D]etermining whether a complaint states a "
        "plausible claim for relief \u2026 [is] a context-specific task that requires the reviewing "
        "court to draw on its judicial experience and common sense.\u201d *Id.* at 679.")
    body_para(doc,
        "Meridian preserved its Rule 12(b)(6) defenses in its Answer\u2014including failure to state a "
        "claim, failure to plead fraud with particularity, the economic-loss doctrine, the DTPA "
        "exemption, the bar on unjust enrichment, and the contractual limitations on liability. "
        "Answer (Dkt. 16) at First Affirmative Defense; id. \u00a7 III (Reservation of Rights). The "
        "Court's Standing Order permits post-answer Rule 12(b)(6) motions where the movant has "
        "preserved the defense, and applies the same plausibility standard. Standing Order \u00a7 3.4. "
        "The same standard governs whether the Motion is analyzed under Rule 12(b)(6) or Rule 12(c). "
        "Claims sounding in fraud must also satisfy the particularity requirements of Rule 9(b). "
        "*Benchmark Elecs., Inc. v. J.M. Huber Corp.*, 343 F.3d 719, 724 (5th Cir. 2003).")

    # IV. ARGUMENT
    heading(doc, "IV.  ARGUMENT", level=1)

    # A. DTPA
    heading(doc, "A.  Count V (DTPA) Must Be Dismissed Because the Transaction Is Exempt Under "
                 "Section 17.49(f).", level=2)
    body_para(doc,
        "Count V fails on a clean, dispositive ground that requires no factual resolution: the "
        "transaction is exempt from the DTPA by statute. Section 17.49(f) of the Texas Business & "
        "Commerce Code exempts from DTPA coverage \u201ca cause of action arising from a transaction, "
        "a project, or a set of transactions relating to the same project, involving total "
        "consideration by the consumer of more than $500,000\u201d unless \u201cthe consumer is an "
        "individual.\u201d Tex. Bus. & Com. Code \u00a7 17.49(f). Both conditions are met here.")
    body_para(doc,
        "First, Arcadia is not an individual. The FAC alleges that Arcadia is \u201ca Texas limited "
        "liability company.\u201d FAC \u00b6 6. A limited liability company is not an \u201cindividual\u201d "
        "within the meaning of Section 17.49(f). Second, the total consideration vastly exceeds "
        "$500,000. The MSLSA alone had a contract value of $14,700,000. FAC \u00b6 34. The four Change "
        "Orders added $2,350,000, for total consideration of $17,050,000. FAC \u00b6\u00b6 46\u201347. "
        "That figure exceeds the statutory threshold by a factor of more than thirty-four. The "
        "exemption is determined by the total consideration involved in the transaction, not by the "
        "amount of damages claimed. *PPG Indus., Inc. v. JMB/Houston Ctrs. Partners Ltd.*, 146 "
        "S.W.3d 79, 89 (Tex. App.\u2014Houston [1st Dist.] 2004, no pet.).")
    body_para(doc,
        "The exemption reflects the Legislature's judgment that sophisticated commercial parties "
        "engaged in high-value transactions should not wield the DTPA\u2014with its treble-damages "
        "provision\u2014as a weapon in ordinary commercial disputes. *Id.* Arcadia is a $62-million "
        "company that negotiated a $14.7 million enterprise-software contract through outside "
        "counsel. FAC \u00b6\u00b6 6, 32. This is precisely the type of transaction the exemption "
        "targets. Count V must be dismissed.")
    body_para(doc,
        "Even absent the exemption, Count V fails on the merits. The alleged representations\u2014"
        "\u201cindustry-leading performance,\u201d \u201c30-40% improvement in reporting efficiency,\u201d "
        "and \u201cseamless\u201d integration\u2014are non-actionable puffery. *Pizza Hut, Inc. v. Papa "
        "John's Int'l, Inc.*, 227 F.3d 489, 497 (5th Cir. 2000). And Arcadia cannot establish "
        "producing cause because the FAC's own allegations and the Change Orders attribute the "
        "implementation difficulties to Arcadia's own delays and to Linden Park's data-migration "
        "errors. *Melody Home Mfg. Co. v. Barnes*, 741 S.W.2d 349, 351\u201352 (Tex. 1987).")

    # B. Unjust Enrichment
    heading(doc, "B.  Count IV (Unjust Enrichment) Must Be Dismissed Because an Express Contract "
                 "Governs the Same Subject Matter.", level=2)
    body_para(doc,
        "Unjust enrichment is unavailable as a matter of law where a valid, express contract governs "
        "the subject matter of the dispute. *Fortune Prod. Co. v. Conoco, Inc.*, 52 S.W.3d 671, 684 "
        "(Tex. 2000). The Texas Supreme Court has held that \u201crecovery under quantum meruit [or "
        "unjust enrichment] is not available when there is a valid, express contract covering the "
        "subject matter of the dispute.\u201d *Id.* The \u201csubject matter\u201d test is applied "
        "broadly: if the contract addresses the general subject area of the dispute, unjust "
        "enrichment is barred. *Excess Underwriters at Lloyd's v. Frank's Casing Crew & Rental "
        "Tools, Inc.*, 246 S.W.3d 42, 59\u201360 (Tex. 2008).")
    body_para(doc,
        "All three conditions are satisfied here. The MSLSA is a valid, express contract that "
        "comprehensively governs the parties' software-licensing and implementation relationship. "
        "FAC \u00b6 69 (alleging \u201ca valid and enforceable written contract\u201d). The subject "
        "matter of Count IV\u2014the $17,050,000 Arcadia paid for software licenses, implementation, "
        "training, support, and data-migration advisory services, FAC \u00b6 100\u2014is identical to "
        "the subject matter of the MSLSA. And Arcadia does not contend that the MSLSA is void or "
        "unenforceable; to the contrary, Arcadia sues on it in Count I. FAC \u00b6 69. Counts I and IV "
        "are therefore irreconcilably inconsistent.")
    body_para(doc,
        "Arcadia's likely response\u2014that the MSLSA's limitation of liability and consequential-"
        "damages waiver prevent full recovery under the contract\u2014is foreclosed. The Texas Supreme "
        "Court has squarely rejected the argument that contractual limitations create a \u201cgap\u201d "
        "that unjust enrichment can fill. *Excess Underwriters*, 246 S.W.3d at 59\u201360. \u201cEquity "
        "follows the law\u201d: where parties have allocated risk through contract, equity will not "
        "rewrite the bargain. *Id.* Count IV must be dismissed.")

    # C. Fraud
    heading(doc, "C.  Count II (Fraud) Must Be Dismissed.", level=2)
    body_para(doc,
        "Count II fails for four independent reasons, any one of which is fatal.")

    heading(doc, "1.  The FAC does not satisfy Rule 9(b).", level=3)
    body_para(doc,
        "Fraud claims must be pleaded with particularity, specifying \u201cthe statements contended to "
        "be fraudulent, identify[ing] the speaker, stat[ing] when and where the statements were "
        "made, and explain[ing] why the statements were fraudulent.\u201d *Benchmark Elecs.*, 343 "
        "F.3d at 724. The FAC identifies the speaker (Poletti) and the approximate dates, but it "
        "does not plead the critical \u201cwhy\u201d: it alleges no facts showing that the "
        "representations were false when made. Instead, the FAC reasons backwards from the "
        "post-Go-Live performance issues to infer that the pre-sale statements must have been "
        "false. FAC \u00b6 80. That post-hoc reasoning is insufficient. *Flaherty & Crumrine "
        "Preferred Income Fund, Inc. v. TXU Corp.*, 565 F.3d 200, 207\u201308 (5th Cir. 2009) "
        "(\u201c[a]llegations of fraud based on post-hoc product or service failures do not, without "
        "more, support an inference that the defendant knew its pre-sale representations were false "
        "at the time they were made\u201d).")
    body_para(doc,
        "The FAC's scienter allegations are equally deficient. Arcadia alleges only that "
        "Meridian's sales team \u201cknew or should have known\u201d the representations were false, "
        "FAC \u00b6 30, and pleads the concealment allegations \u201con information and belief.\u201d "
        "FAC \u00b6\u00b6 56\u201360. While Rule 9(b) permits scienter to be alleged generally, the "
        "plaintiff must still plead facts giving rise to a strong inference of fraudulent intent. "
        "*Flaherty*, 565 F.3d at 207. The FAC alleges no internal documents, no contemporaneous "
        "admissions, and no specific facts establishing knowledge of falsity at the time of the "
        "representations. The FAC also improperly \u201clumps\u201d multiple pre-sale communications "
        "together without individually analyzing each statement. *Dorsey v. Portfolio Equities, "
        "Inc.*, 540 F.3d 333, 340 (5th Cir. 2008).")

    heading(doc, "2.  The alleged representations are non-actionable puffery.", level=3)
    body_para(doc,
        "The core of Arcadia's fraud claim rests on three statements: that NexusCore \u201cdelivers "
        "industry-leading performance for healthcare analytics,\u201d that clients \u201ctypically see "
        "30-40% improvement in reporting efficiency,\u201d and that NexusCore \u201cwill integrate "
        "seamlessly with any EHR platform.\u201d FAC \u00b6\u00b6 23, 25, 79. These are textbook puffery. "
        "The Fifth Circuit holds that \u201ca general claim of superiority that is so vague that it "
        "can be understood only as the seller's opinion of the product rather than a specific, "
        "factual representation\u201d is non-actionable. *Pizza Hut*, 227 F.3d at 497; see also "
        "*Castrol Inc. v. Pennzoil Co.*, 987 F.2d 939, 945 (3d Cir. 1993) (defining puffery as "
        "\u201cexaggerated advertising, blustering, and boasting upon which no reasonable buyer would "
        "rely\u201d). \u201cIndustry-leading performance\u201d is a vague, subjective claim of product "
        "superiority. \u201c30-40% improvement\u201d is a generalized projection of typical results, "
        "not a binding guarantee. And \u201cseamless\u201d integration is a general claim of "
        "capability, not a specific factual representation.")
    body_para(doc,
        "The puffery remains non-actionable even though it appeared in a sales presentation that "
        "also contained factual content. *Pizza Hut*, 227 F.3d at 498\u201399. And forward-looking "
        "statements about product performance made in a sales context by a party with an obvious "
        "incentive to promote the product are puffery upon which a sophisticated buyer cannot "
        "reasonably rely. *Presidio Enters., Inc. v. Warner Bros. Distrib. Corp.*, 784 F.2d 674, "
        "679\u201380 (5th Cir. 1986). The December 9, 2021 written proposal\u2014which the FAC "
        "acknowledges contained \u201ccertain general disclaimers,\u201d FAC \u00b6 27\u2014expressly "
        "stated that \u201cestimated timelines and performance metrics are provided for planning "
        "purposes only and do not constitute guarantees\u201d and that \u201cMeridian does not warrant "
        "that integration with any particular EHR system will be seamless, error-free, or "
        "achievable within any specific timeframe.\u201d")

    heading(doc, "3.  The integration clause and anti-reliance doctrine bar reliance on "
                 "extra-contractual representations.", level=3)
    body_para(doc,
        "Section 12.1 of the MSLSA is a comprehensive integration clause that supersedes \u201call "
        "prior and contemporaneous agreements, proposals, representations, and understandings, "
        "whether oral or written,\u201d including \u201cmarketing materials, product descriptions, "
        "brochures, presentations, and demonstrations.\u201d MSLSA \u00a7 12.1. The MSLSA is governed "
        "by Delaware law. MSLSA \u00a7 12.7. Under Delaware law, \u201ca party to a contract cannot "
        "promise, in a clear integration clause of a negotiated agreement, that it is not relying "
        "on promises or representations made by the other party outside of the agreement, and then "
        "assert a claim for fraud based on those very representations.\u201d *Eagle Indus., Inc. v. "
        "DeVilbiss Health Care, Inc.*, 702 A.2d 1228, 1232 (Del. 1997); see also *SIGA Techs., Inc. "
        "v. PharmAthene, Inc.*, 67 A.3d 330, 344 (Del. 2013). The FAC's fraud claim rests entirely "
        "on extra-contractual representations\u2014Poletti's sales presentations and the pre-contract "
        "demonstration\u2014precisely the representations Section 12.1 was designed to supersede.")
    body_para(doc,
        "This result is compelled by the FAC's own admissions. Arcadia was \u201crepresented by "
        "counsel\u201d and \u201cma[d]e efforts to negotiate more favorable terms, including enhanced "
        "limitation of liability provisions and expanded warranty protections.\u201d FAC \u00b6 32. "
        "Sophisticated parties represented by counsel are presumed to have understood the "
        "integration clause. *Kuhn Constr., Inc. v. Diamond State Port Corp.*, 990 A.2d 393, 401 "
        "(Del. 2010). Arcadia cannot now claim reliance on pre-contractual representations that the "
        "integration clause\u2014to which Arcadia, with the benefit of counsel, agreed\u2014was designed "
        "to bar.")

    heading(doc, "4.  The economic-loss doctrine and the absence of justifiable reliance bar the "
                 "fraud claim.", level=3)
    body_para(doc,
        "All of Arcadia's damages are purely economic, and the alleged misrepresentations concern "
        "the quality and performance of NexusCore\u2014the subject matter of the MSLSA. Under "
        "Delaware law, the economic-loss doctrine bars tort claims that seek to recover economic "
        "losses arising from a contractual relationship absent an independent duty of care. *Brasby "
        "v. Morris Dynamics, Inc.*, 947 A.2d 1042, 1049 (Del. 2008); *Kuhn Constr.*, 990 A.2d at "
        "401\u2013402. Under Texas law, the same result obtains where (1) the plaintiff's losses are "
        "purely economic, (2) the losses relate to the subject matter of a contract, and (3) the "
        "defendant owes no duty independent of the contract. *Chapman Custom Homes, Inc. v. Dallas "
        "Plumbing Co.*, 445 S.W.3d 716, 718\u201319 (Tex. 2014); *Sharyland Water Supply Corp. v. City "
        "of Alton*, 354 S.W.3d 407, 415\u201316 (Tex. 2011). All three factors are met. Arcadia seeks "
        "the contract value, lost profits, increased operating costs, and reputational harm\u2014all "
        "economic losses tied to NexusCore. FAC \u00b6\u00b6 61\u201366. Meridian owes no duty to "
        "Arcadia independent of the MSLSA.")
    body_para(doc,
        "Justifiable reliance is also absent as a matter of law. The FAC admits that Arcadia "
        "\u201cconducted certain independent due diligence, including reference calls with two "
        "existing Meridian healthcare clients\u201d before contracting. FAC \u00b6 31. A plaintiff who "
        "conducts an independent investigation and forms its own conclusions cannot later claim it "
        "was defrauded by the seller's promotional statements. *Presidio Enters.*, 784 F.2d at "
        "680. Combined with the proposal disclaimers, the integration clause, and Arcadia's "
        "representation by counsel, any reliance on general sales statements was not justifiable. "
        "Count II must be dismissed.")

    # D. Negligent Misrepresentation
    heading(doc, "D.  Count III (Negligent Misrepresentation) Must Be Dismissed.", level=2)
    body_para(doc,
        "Count III fails for three independent reasons. First, Arcadia cannot identify a duty "
        "independent of the MSLSA. Texas recognizes negligent misrepresentation under Restatement "
        "(Second) of Torts \u00a7 552, but liability is constrained by the independent-duty "
        "requirement. *McCamish, Martin, Brown & Loeffler v. F.E. Appling Interests*, 991 S.W.2d "
        "787, 791\u201392 (Tex. 1999). Meridian's relationship with Arcadia is purely contractual; "
        "there is no fiduciary or professional relationship giving rise to heightened duties. The "
        "MSLSA's warranty provisions (Section 9.1) and disclaimers (Section 9.4) define the scope "
        "of Meridian's obligations regarding software performance. Arcadia cannot circumvent those "
        "provisions through a negligent-misrepresentation claim.")
    body_para(doc,
        "Second, the economic-loss doctrine bars the claim for the same reasons discussed above. "
        "*Chapman Custom Homes*, 445 S.W.3d at 718\u201319; *Brasby*, 947 A.2d at 1049. Third, "
        "Arcadia's damages are benefit-of-the-bargain losses\u2014the contract value, lost profits, "
        "and consequential costs\u2014which are recoverable only in contract, not in tort. *Federal "
        "Land Bank Ass'n v. Sloane*, 825 S.W.2d 439, 442\u201343 (Tex. 1992). Where the plaintiff's "
        "true complaint is that the other party did not deliver what was contractually promised, "
        "the claim sounds in contract. *Id.* at 443. Finally, because Count III sounds in fraud, it "
        "must satisfy Rule 9(b), and for the reasons stated in Section IV.C.1, it does not. "
        "*Benchmark Elecs.*, 343 F.3d at 724. Count III must be dismissed.")

    # E. Breach of Contract
    heading(doc, "E.  Count I (Breach of Contract) Must Be Dismissed, or in the Alternative, "
                 "Arcadia's Damages Must Be Limited to the Contractual Cap.", level=2)
    body_para(doc,
        "Count I is defeated by the MSLSA's deemed-acceptance provision, the expired limited "
        "warranty, and the contractual limitations on remedies and damages. The MSLSA is governed "
        "by Delaware law. MSLSA \u00a7 12.7. Delaware courts strictly enforce negotiated "
        "risk-allocation provisions between sophisticated parties. *Kana Software, Inc. v. Sealand "
        "Tech., Inc.*, 178 A.3d 1045, 1058\u201360 (Del. Ch. 2017); *Dresser-Rand Co. v. Virtual "
        "Automation, Inc.*, 361 F.3d 831, 838\u201340 (5th Cir. 2004) (enforcing liability cap and "
        "consequential-damages waiver in negotiated software agreement under both Texas and "
        "Delaware law).")

    heading(doc, "1.  Arcadia's breach-of-contract claim is barred by deemed acceptance.", level=3)
    body_para(doc,
        "Section 5.3 gave Arcadia thirty days following Go-Live to conduct acceptance testing and "
        "deliver written notice of any material nonconformity; absent such notice, the software is "
        "\u201cdeemed accepted.\u201d MSLSA \u00a7 5.3(d). Go-Live occurred on January 15, 2023, so the "
        "acceptance window closed on February 14, 2023. Arcadia did not deliver any written notice "
        "of nonconformity within that period. Dr. Okonkwo's March 8, 2023 email\u2014which reported "
        "only \u201cintermittent latency issues\u201d and \u201coccasional report generation errors\u201d "
        "and noted that no patient-safety concerns existed, FAC \u00b6 50\u2014arrived twenty-two days "
        "after the window closed. Under Section 5.3, acceptance was conclusively established on "
        "February 14, 2023.")
    body_para(doc,
        "Deemed-acceptance provisions are enforceable and provide a bright-line rule. *Simulados, "
        "Inc. v. Canton Health Mgmt. Co.*, 2019 WL 4573218 (W.D. Tex. Sept. 20, 2019) "
        "(unpublished) (enforcing 30-day deemed-acceptance provision; oral complaints do not "
        "satisfy written-notice requirement); *Precision Healthcare Sols. v. Nextera Data Sys.*, "
        "458 F. Supp. 3d 544, 551\u201352 (N.D. Tex. 2020) (deemed acceptance applies to both patent "
        "and latent defects; risk of inadequate testing falls on the licensee). Once acceptance is "
        "deemed granted, Arcadia's remedies for alleged defects are limited to the warranty "
        "provisions of Section 9.1, not pre-acceptance rejection rights. *Id.* at 552. The FAC "
        "alleges no facts showing that Meridian intentionally concealed known defects to prevent "
        "discovery during the testing period. *Id.*")

    heading(doc, "2.  The limited warranty has expired, and its remedy is exclusive.", level=3)
    body_para(doc,
        "Section 9.1 provides a limited warranty that the software will \u201cperform substantially in "
        "accordance with the Documentation\u201d for ninety (90) days following Go-Live. MSLSA \u00a7 "
        "9.1(a). That period expired on April 15, 2023. The sole and exclusive remedy for any "
        "breach of that warranty is repair, replacement, or\u2014if Meridian cannot cure within sixty "
        "days\u2014a pro-rata refund of pre-paid license fees. MSLSA \u00a7 9.1(b). Section 8.3 "
        "designates the MSLSA's remedies as \u201csole and exclusive.\u201d MSLSA \u00a7 8.3. Arcadia's "
        "demand for $47.3 million\u2014including the full contract value as \u201crestitution,\u201d FAC "
        "\u00b6 62\u2014far exceeds the exclusive contractual remedy and is therefore not recoverable. "
        "*Kana Software*, 178 A.3d at 1059 (where contract provides a specific remedy for "
        "nonconformance, that remedy is exclusive and precludes tort-based damage theories).")

    heading(doc, "3.  The consequential-damages waiver and liability cap bar the bulk of "
                 "Arcadia's damages.", level=3)
    body_para(doc,
        "Even if some breach-of-contract theory were to survive, the vast majority of Arcadia's "
        "claimed damages are contractually barred. Section 8.2 excludes all \u201cindirect, "
        "incidental, special, consequential, punitive, or exemplary damages, including \u2026 damages "
        "for loss of profits, goodwill, use, data, or other intangible losses.\u201d MSLSA \u00a7 8.2. "
        "Arcadia's $18.4 million lost-profits claim, $6.7 million increased-operating-costs claim, "
        "and $5.0 million reputational-harm claim are consequential damages expressly waived. FAC "
        "\u00b6\u00b6 63\u201365. *Dresser-Rand*, 361 F.3d at 839\u201340 (enforcing consequential-"
        "damages waiver as independent provision). Section 8.1 further caps aggregate liability at "
        "the fees paid or payable in the twelve months preceding the event giving rise to "
        "liability. MSLSA \u00a7 8.1. These provisions reflect a bargained-for allocation of risk "
        "that the FAC admits Arcadia negotiated. FAC \u00b6 32. *Dresser-Rand*, 361 F.3d at 839.")
    body_para(doc,
        "Finally, the data-migration defects that pervade the FAC are not attributable to Meridian. "
        "The MSLSA and SOW-1 allocate all data-migration responsibility to Arcadia. SOW-1 \u00a7 3.2; "
        "MSLSA Ex. C \u00a7 C.2. And CO-004\u2014signed by Arcadia\u2014admits that the data-integrity "
        "issues \u201carose from errors in data migration scripts prepared and executed by Linden Park "
        "Consulting, LLC, and not from any defect in the NexusCore software or any act or omission "
        "of [Meridian].\u201d CO-004 \u00a7 5.1. The FAC's contrary allegations, FAC \u00b6\u00b6 43, 52, "
        "are directly contradicted by the contract and the Change Order and cannot survive. Count I "
        "must be dismissed, or in the alternative, Arcadia's damages must be limited to the "
        "contractual cap and the exclusive warranty remedy.")

    # V. CONCLUSION
    heading(doc, "V.  CONCLUSION", level=1)
    body_para(doc,
        "For the foregoing reasons, Defendant Meridian Cloud Solutions, Inc. respectfully requests "
        "that this Court enter an Order dismissing all five counts of Plaintiff's First Amended "
        "Complaint with prejudice, and awarding Meridian such other and further relief as the Court "
        "deems just and proper. A proposed order accompanies this Motion.")

    # Signature
    signature_block(doc, MERIDIAN_COUNSEL,
                    "Dated: January 15, 2024",
                    FIRM_LINES, PARTY_LINE)

    # Certificate of Service
    heading(doc, "CERTIFICATE OF SERVICE", level=1, space_before=14)
    body_para(doc,
        "I hereby certify that on January 15, 2024, I electronically filed the foregoing Defendant "
        "Meridian Cloud Solutions, Inc.'s Motion to Dismiss Plaintiff's First Amended Complaint "
        "Pursuant to Federal Rule of Civil Procedure 12(b)(6) with the Clerk of Court using the "
        "CM/ECF system, which will send notification of such filing to all counsel of record, "
        "including:")
    p = doc.add_paragraph(); _para_format(p, line=2.0, left_indent=0.5, space_after=0)
    add_rich(p, f"{PLAINTIFF_COUNSEL[0]}")
    p = doc.add_paragraph(); _para_format(p, line=2.0, left_indent=0.5, space_after=0)
    add_rich(p, PLAINTIFF_COUNSEL[1])
    p = doc.add_paragraph(); _para_format(p, line=2.0, left_indent=0.5, space_after=0)
    add_rich(p, PLAINTIFF_COUNSEL[2])
    p = doc.add_paragraph(); _para_format(p, line=2.0, left_indent=0.5, space_after=6)
    add_rich(p, f"Email: {PLAINTIFF_COUNSEL[3]}")
    p = doc.add_paragraph(); _para_format(p, line=2.0, left_indent=0.5, space_after=0)
    add_rich(p, "Counsel for Plaintiff Arcadia Health Systems, LLC", base_italic=True)
    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=6)
    p = doc.add_paragraph(); _para_format(p, line=2.0, space_after=0)
    add_rich(p, "/s/ Margaret \u201cMeg\u201d Calloway")
    p = doc.add_paragraph(); _para_format(p, line=2.0, space_after=0)
    add_rich(p, "Margaret \u201cMeg\u201d Calloway")

    doc.save("output/motion-to-dismiss.docx")
    print("Wrote output/motion-to-dismiss.docx")


# ============================================================================
# DOCUMENT 2: PROPOSED ORDER
# ============================================================================

def build_order():
    doc = Document()
    setup_document(doc)
    caption(doc, "PROPOSED ORDER GRANTING DEFENDANT'S MOTION TO DISMISS")

    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=2.0, space_after=6)
    add_rich(p,
        "On this day, the Court considered Defendant Meridian Cloud Solutions, Inc.'s Motion to "
        "Dismiss Plaintiff's First Amended Complaint Pursuant to Federal Rule of Civil Procedure "
        "12(b)(6) (the \u201cMotion\u201d), together with the accompanying brief and proposed order. "
        "Plaintiff Arcadia Health Systems, LLC filed a response in opposition, and Defendant filed "
        "a reply. Having reviewed the Motion, the responsive briefing, the First Amended Complaint "
        "(Dkt. 12), the Answer (Dkt. 16), and the applicable law, and being fully advised, the "
        "Court is of the opinion that the Motion should be and hereby is GRANTED.")

    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=2.0, space_after=6)
    add_rich(p,
        "It is therefore ORDERED, ADJUDGED, and DECREED that:")

    items = [
        "Defendant's Motion to Dismiss (Dkt. [__]) is GRANTED.",
        "Count I (Breach of Contract) is DISMISSED WITH PREJUDICE.",
        "Count II (Fraud) is DISMISSED WITH PREJUDICE.",
        "Count III (Negligent Misrepresentation) is DISMISSED WITH PREJUDICE.",
        "Count IV (Unjust Enrichment) is DISMISSED WITH PREJUDICE.",
        "Count V (Violation of the Texas Deceptive Trade Practices\u2013Consumer Protection Act) is "
        "DISMISSED WITH PREJUDICE.",
        "All claims asserted by Plaintiff Arcadia Health Systems, LLC against Defendant Meridian "
        "Cloud Solutions, Inc. in the First Amended Complaint are DISMISSED WITH PREJUDICE.",
        "The Clerk of Court is DIRECTED to enter final judgment in favor of Defendant Meridian "
        "Cloud Solutions, Inc. and against Plaintiff Arcadia Health Systems, LLC, with each party "
        "to bear its own costs, except as otherwise provided by law or contract.",
        "All pending motions, if any, are DENIED AS MOOT.",
    ]
    for it in items:
        p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=2.0,
                                              left_indent=0.5, first_line=-0.25, space_after=4)
        add_rich(p, it)

    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=6)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=2.0, space_after=6)
    add_rich(p,
        "This is a final, appealable order. The Court finds that there is no just reason for delay "
        "and directs the entry of final judgment pursuant to Federal Rule of Civil Procedure 54(b).")

    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=10)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=2.0, space_after=0)
    add_rich(p, "SIGNED and ENTERED on this _____ day of ______________, 2024.")
    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=18)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
    add_rich(p, "____________________________________")
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
    add_rich(p, "THE HONORABLE CATHERINE M. ALVAREZ", base_bold=True)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
    add_rich(p, "United States District Judge")
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.15, space_after=0)
    add_rich(p, "Western District of Texas, Austin Division")

    doc.save("output/proposed-order.docx")
    print("Wrote output/proposed-order.docx")


# ============================================================================
# DOCUMENT 3: COVER MEMO
# ============================================================================

def build_memo():
    doc = Document()
    setup_document(doc)

    # Memo header
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, space_after=0)
    add_rich(p, "STONEBRIDGE & CALLOWAY LLP", base_bold=True)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, space_after=2)
    add_rich(p, "600 Congress Avenue, Suite 2800  \u2022  Austin, Texas 78701", size=11)
    hr = doc.add_paragraph(); _para_format(hr, line=1.0, space_after=6)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom); pPr.append(pbdr)

    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, space_after=2)
    add_rich(p, "PRIVILEGED & CONFIDENTIAL  \u2014  ATTORNEY WORK PRODUCT", base_bold=True, size=11)
    p = doc.add_paragraph(); _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15, space_after=8)
    add_rich(p, "INTERNAL MEMORANDUM", base_bold=True)

    # TO/FROM/DATE/RE block
    fields = [
        ("TO:", "Margaret \u201cMeg\u201d Calloway, Lead Partner; David Arsenault, Senior Associate"),
        ("FROM:", "Litigation Team, Stonebridge & Calloway LLP"),
        ("DATE:", "January 12, 2024"),
        ("RE:", "Threshold Issues \u2014 Rule 12(b)(6) Motion to Dismiss, *Arcadia Health Systems, "
                "LLC v. Meridian Cloud Solutions, Inc.*, No. 1:23-cv-00847-CMA (W.D. Tex. Austin "
                "Division)"),
    ]
    for label, val in fields:
        p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=2)
        r = p.add_run(label + " "); _set_run_font(r, bold=True)
        add_rich(p, val)
    hr = doc.add_paragraph(); _para_format(hr, line=1.0, space_after=8)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom); pPr.append(pbdr)

    body_para(doc,
        "This memorandum flags threshold issues that the team should resolve before filing the "
        "Rule 12(b)(6) motion to dismiss (the \u201cMotion\u201d), currently scheduled for filing on "
        "January 15, 2024. The Motion, proposed order, and this memo are attached. The issues "
        "below are listed in order of urgency. The first\u2014a potential subject-matter-jurisdiction "
        "defect\u2014is dispositive and may require a change in strategy before we file anything on "
        "the merits.")

    # 1. Jurisdiction
    heading(doc, "1.  Subject-Matter Jurisdiction \u2014 A Likely Complete-Diversity Defect (HIGHEST PRIORITY)", level=1)
    body_para(doc,
        "Our Notice of Removal asserted diversity jurisdiction under 28 U.S.C. \u00a7 1332(a) but "
        "analyzed Arcadia's citizenship based solely on its managing member, Dr. Okonkwo (a Texas "
        "citizen). Notice of Removal \u00b6\u00b6 13\u201314. That analysis is incomplete. An LLC's "
        "citizenship is determined by the citizenship of *all* of its members. *Harvey v. Grey Wolf "
        "Drilling Co.*, 542 F.3d 1077, 1080 (5th Cir. 2008). The FAC and Arcadia's Operating "
        "Agreement disclose three members: (a) Dr. Rachel Okonkwo (Texas); (b) Martin Schreiber "
        "(Texas); and (c) Apex Medical Ventures, LP, a Delaware limited partnership holding a 30% "
        "interest. FAC \u00b6\u00b6 7\u20138; Operating Agreement Ex. A.")
    body_para(doc,
        "Apex's citizenship, in turn, is determined by the citizenship of its partners. *Carden v. "
        "Arkoma Assocs., Ltd.*, 494 U.S. 185, 195\u201396 (1990). Apex's general partner is Apex "
        "Medical Ventures GP, Inc., a Delaware corporation with its principal place of business in "
        "Wilmington, Delaware\u2014making it a citizen of Delaware. Operating Agreement Ex. A, Note "
        "3(a). Meridian is also a citizen of Delaware (incorporation) and Texas (principal place of "
        "business, Austin). FAC \u00b6 9; Notice of Removal \u00b6\u00b6 9\u201311. Because both Arcadia "
        "(through Apex's Delaware general partner) and Meridian are citizens of Delaware, complete "
        "diversity is destroyed. *Strawbridge v. Curtiss*, 7 U.S. (3 Cranch) 267 (1806).")
    body_para(doc,
        "Subject-matter jurisdiction cannot be waived and may be raised at any time, by a party or "
        "sua sponte by the Court. *Arbaugh v. Y & H Corp.*, 546 U.S. 500, 514 (2006). If the "
        "defect is confirmed, the Court lacks jurisdiction and the case must be remanded to Bexar "
        "County state court. This could render the Rule 12(b)(6) Motion moot.")
    body_para(doc,
        "Recommended action: (i) Confirm the citizenship analysis by obtaining the Apex Medical "
        "Ventures, LP partnership agreement (we have the Operating Agreement, which is sufficient "
        "to confirm the GP's Delaware citizenship, but the partnership agreement would lock it "
        "down); (ii) Decide strategically whether Meridian is better off in federal court (Judge "
        "Alvarez's rigorous pleading standards, bench trial, Delaware choice of law) or state "
        "court (Bexar County, possible jury). If we conclude the diversity defect is real, we "
        "should consider moving to remand *before* filing the 12(b)(6) Motion\u2014or at minimum be "
        "prepared for the Court to raise jurisdiction sua sponte. Filing a merits motion in a court "
        "that may lack jurisdiction is not ideal, but the jurisdictional issue is independent of "
        "the merits and can be raised separately. We should flag this for the client immediately.")

    # 2. Post-answer motion
    heading(doc, "2.  Post-Answer Motion \u2014 Rule 12(b)(6) vs. Rule 12(c)", level=1)
    body_para(doc,
        "Meridian already filed its Answer (Oct. 2, 2023) and preserved the Rule 12(b)(6) defense "
        "(First Affirmative Defense; \u00a7 III, Reservation of Rights; Nov. 8, 2023 notice letter). "
        "The Standing Order (\u00a7 3.4) permits post-answer Rule 12(b)(6) motions where the defense "
        "is preserved and says the Court will treat them as Rule 12(b)(6) motions. The Scheduling "
        "Order, however, states the Court \u201cwill treat any such motion as a Rule 12(c) motion for "
        "judgment on the pleadings, evaluated under the same standard as a Rule 12(b)(6) motion.\u201d "
        "There is a tension, but it is immaterial: the plausibility standard is identical under "
        "either rule, and the same documents may be considered. The Motion is captioned as a "
        "Rule 12(b)(6) motion per the Standing Order and the task; the standard section notes that "
        "the same analysis applies under Rule 12(c). No change needed, but be prepared to address "
        "the characterization if the Court raises it.")

    # 3. Choice of law
    heading(doc, "3.  Choice of Law \u2014 Delaware (contract) vs. Texas (tort/DTPA)", level=1)
    body_para(doc,
        "Section 12.7 of the MSLSA selects Delaware law (without conflicts principles) to govern "
        "\u201cthis Agreement.\u201d Count I (breach of contract) is governed by Delaware law, and the "
        "Motion relies on Delaware authority (*Kana Software*, *Brasby*, *Kuhn*, *Eagle Industries*, "
        "*SIGA*) for the contractual limitations, deemed acceptance, and economic-loss arguments. "
        "For the tort claims (Counts II\u2013III), the scope of \u00a7 12.7 is debatable; Texas courts "
        "apply the most-significant-relationship test, and all contacts point to Texas. The Motion "
        "therefore argues the tort claims fail under *both* Delaware and Texas law. The DTPA "
        "(Count V) is a Texas statute regardless. Unjust enrichment (Count IV) yields the same "
        "result under either body of law. We should be ready to brief choice of law more fully if "
        "Arcadia contests it, but the Motion's dual-track approach covers both outcomes.")

    # 4. DTPA pre-suit notice
    heading(doc, "4.  DTPA Pre-Suit Notice \u2014 Possible \u00a7 17.505(a) Timing Defect", level=1)
    body_para(doc,
        "Arcadia's DTPA pre-suit demand letter is dated June 15, 2023; suit was filed August 7, "
        "2023\u2014approximately 53 days later. Section 17.505(a) requires the consumer to give "
        "written notice \u201cat least 60 days before filing the suit.\u201d If the demand was received "
        "in mid-June, Arcadia may have filed before the 60-day period elapsed. (Note: Meridian's "
        "July 10, 2023 response letter references a demand \u201cdated June 26, 2023,\u201d which would "
        "make the shortfall even larger; the team should reconcile the dates.) The remedy for a "
        "\u00a7 17.505(a) violation is typically abatement rather than dismissal, so this is a "
        "secondary issue. The Motion leads with the \u00a7 17.49(f) exemption, which is dispositive "
        "and does not depend on the notice timing. We can raise the \u00a7 17.505(a) defect "
        "separately or in a reply if useful, but it should not displace the exemption argument.")

    # 5. Count I strategy
    heading(doc, "5.  Count I Strategy \u2014 Full Dismissal vs. Damages Limitation", level=1)
    body_para(doc,
        "The Motion takes the aggressive position that Count I should be dismissed outright based "
        "on deemed acceptance (\u00a7 5.3), the expired limited warranty (\u00a7 9.1), and the exclusive "
        "remedies provision (\u00a7 8.3), with the alternative that damages be limited to the "
        "contractual cap (\u00a7\u00a7 8.1, 8.2). Full dismissal of a breach-of-contract claim at the "
        "12(b)(6) stage is aggressive; courts are sometimes reluctant to dismiss contract claims "
        "where the deemed-acceptance and warranty arguments turn on the timing of notice. Here, "
        "however, the dates are undisputed on the face of the FAC (Go-Live Jan. 15; first written "
        "complaint Mar. 8) and the contract (30-day window; 90-day warranty), so the arguments are "
        "well-suited to resolution on the pleadings. The alternative damages-limitation argument "
        "preserves a strong fallback: even if the Court allows the contract claim to proceed, the "
        "$18.4M lost-profits, $6.7M operating-costs, and $5.0M reputational-harm components are "
        "consequential damages barred by \u00a7 8.2, and the remainder is capped by \u00a7 8.1. We "
        "should confirm the 12-month fee figure for the \u00a7 8.1 cap (see item 8 below).")

    # 6. With/without prejudice
    heading(doc, "6.  Dismissal With vs. Without Prejudice", level=1)
    body_para(doc,
        "The proposed order dismisses all counts with prejudice, consistent with the Answer's "
        "prayer. With prejudice is clearly appropriate for Count V (the \u00a7 17.49(f) exemption is "
        "a legal bar no amendment can cure) and Count IV (unjust enrichment is unavailable as a "
        "matter of law while the MSLSA governs). For Counts II and III, dismissal with prejudice "
        "is defensible because the defects are legal (puffery, economic-loss doctrine, integration "
        "clause, no independent duty) and not curable by additional facts, but the Court may grant "
        "leave to amend. For Count I, the Court may be inclined to allow amendment. We should be "
        "prepared to oppose any motion for leave to amend by arguing futility, but should not be "
        "surprised if the Court dismisses some counts without prejudice. Consider whether to "
        "request dismissal with prejudice in the Motion's prayer while acknowledging the Court's "
        "discretion\u2014the current draft does so.")

    # 7. Documents on the motion
    heading(doc, "7.  Scope of Materials Considered on the Motion \u2014 Avoid Conversion", level=1)
    body_para(doc,
        "The Motion relies only on the FAC and documents referenced in and central to it: the "
        "MSLSA, SOW-1, the four Change Orders (including CO-004's admission that Linden Park caused "
        "the data issues), and the December 9, 2021 proposal (whose disclaimers the FAC itself "
        "acknowledges). This keeps the Motion within the Rule 12(b)(6) framework and avoids "
        "conversion to summary judgment under Rule 12(d). *Lormand*, 565 F.3d at 255\u201356. "
        "Critically, the Motion does *not* rely on materials outside the pleadings\u2014namely: "
        "Schreiber's January 19, 2022 internal due-diligence memo (which acknowledges integration "
        "complexity and recommends an SI partner), Poletti's contemporaneous notes of the Nov. 18, "
        "2021 meeting, the support-ticket analysis (47 tickets; 31 in SLA; 12 traceable to Linden "
        "Park; 4 patched), and the Aug. 25, 2022 status report's root-cause findings (45-day PM "
        "gap; 75-day API-spec delay). These are powerful but extrinsic. They should be preserved "
        "and deployed at the summary-judgment stage, not in this Motion. Do not add them to the "
        "Motion or attach them as exhibits now.")

    # 8. Open items
    heading(doc, "8.  Open Items to Resolve Before Filing", level=1)
    body_para(doc,
        "(a) \u00a7 8.1 liability cap amount. The cap is \u201ctotal fees paid or payable by Licensee "
        "during the twelve (12) month period immediately preceding the event giving rise to such "
        "liability.\u201d We need Meridian's billing records to compute the exact figure, and we must "
        "identify the \u201cevent giving rise to liability\u201d (the pre-sale representations in "
        "Oct.\u2013Nov. 2021, the Jan. 15, 2023 Go-Live, or the Mar. 8, 2023 first complaint). This "
        "affects the alternative damages-limitation argument. The Motion references the cap "
        "generally; we can supply the figure in the reply if needed.")
    body_para(doc,
        "(b) \u00a7 17.49(f) alternative \u201cassets\u201d basis. Section 17.49(f) also exempts "
        "transactions where the consumer's assets exceed $25 million. We requested Arcadia's "
        "financial documentation in our July 10, 2023 DTPA-response letter. If Arcadia's assets "
        "exceeded $25 million as of March 15, 2022 (likely, given ~$62M revenue and 23-hospital "
        "operations), this is an additional, independent basis for the exemption. Pursue this in "
        "discovery if the Court does not dismiss Count V on the $500,000 ground alone.")
    body_para(doc,
        "(c) Certificate of conference. The Standing Order (\u00a7 3.1) requires a certificate of "
        "conference and states it \u201cshall be filed as a separate document concurrently with the "
        "motion.\u201d The Motion includes a certificate-of-conference section, but per the Standing "
        "Order we should also file it as a separate docket entry. Complete the meet-and-confer with "
        "Breckenridge (we have the Nov. 8 letter and subsequent conferrals) and file the separate "
        "certificate alongside the Motion.")
    body_para(doc,
        "(d) Proposed order docket number. The Standing Order (\u00a7 3.5) requires the proposed "
        "order to identify the motion by docket entry number; if unknown at filing, counsel must "
        "update the proposed order with the correct number within one business day. The proposed "
        "order uses \u201cDkt. [__]\u201d placeholders; update them once the Motion is docketed.")
    body_para(doc,
        "(e) Page limits and formatting. The Motion is within the 25-page limit (exclusive of "
        "caption, signature block, and certificates). Confirm the final page count after rendering. "
        "Reply briefs are capped at 10 pages under the Scheduling Order (the Standing Order says "
        "15; the Scheduling Order's 10-page figure controls as the case-specific order). Plan the "
        "reply accordingly.")
    body_para(doc,
        "(f) Jury-trial waiver. Section 12.9 of the MSLSA waives jury trial; the Scheduling Order "
        "sets a bench trial. The FAC demands a jury and our Answer conditionally demands one. This "
        "is a trial-stage issue, not a threshold issue for the Motion, but note that the "
        "enforceability of the jury waiver under Delaware law (which governs the MSLSA) may be "
        "contested. No action needed for the Motion.")

    # 9. Bottom line
    heading(doc, "9.  Bottom Line", level=1)
    body_para(doc,
        "The Motion presents strong, well-supported arguments for dismissal of all five counts. "
        "The DTPA exemption (\u00a7 17.49(f)) and the unjust-enrichment bar are clean legal rules that "
        "should result in dismissal. The fraud and negligent-misrepresentation claims fail under "
        "Rule 9(b), the puffery doctrine, the integration clause, and the economic-loss doctrine. "
        "The breach-of-contract claim is defeated by deemed acceptance and the expired exclusive "
        "warranty, with a robust alternative damages-limitation argument. The single most "
        "important item to resolve before filing is the diversity-jurisdiction defect (item 1), "
        "which is independent of the merits but could change where and how we litigate this case. "
        "Please advise on the jurisdictional strategy at your earliest convenience.")

    sp = doc.add_paragraph(); _para_format(sp, line=1.0, space_after=6)
    p = doc.add_paragraph(); _para_format(p, line=1.15, space_after=0)
    add_rich(p, "Attachments: (1) Motion to Dismiss; (2) Proposed Order.")

    doc.save("output/cover-memo.docx")
    print("Wrote output/cover-memo.docx")


if __name__ == "__main__":
    # Two-pass for the motion: build once with placeholder TOA pages, compute the
    # actual page numbers from a rendered PDF, then rebuild with real page numbers.
    build_motion({})
    toa_pages = compute_toa_pages("output/motion-to-dismiss.docx")
    build_motion(toa_pages)
    build_order()
    build_memo()
    print("All documents generated.")
    print("TOA pages:", toa_pages)